# snt/core.py
import pandas as pd
from pathlib import Path
import numpy as np
from tabulate import tabulate

from pathlib import Path
import pandas as pd
from tabulate import tabulate
import math

#
def create_comprehensive_trend_maps(output_dir='trend_maps/'):
    import os
    import matplotlib.pyplot as plt
    import pandas as pd
    import numpy as np
    import geopandas as gpd
    import re
    from matplotlib.colors import BoundaryNorm
    from matplotlib.cm import RdBu_r
    import matplotlib.patches as mpatches
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Read data
    df1 = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    
    # Identify crude incidence columns for 2021-2024
    pattern = re.compile(r'^crude_incidence_(\d{4})$')
    year_cols = [col for col in df1.columns 
                 if pattern.match(col) and 2021 <= int(pattern.match(col).group(1)) <= 2024]
    
    if len(year_cols) < 2:
        print("Error: Need at least 2 years of data to calculate percentage change")
        return
    
    # Calculate overall percentage change for each chiefdom (2021 to 2024 specifically)
    def calculate_overall_change(row):
        # Look specifically for 2021 and 2024 data
        value_2021 = None
        value_2024 = None
        
        for col in year_cols:
            year = int(re.search(r'(\d{4})', col).group(1))
            if year == 2021 and pd.notna(row[col]):
                value_2021 = row[col]
            elif year == 2024 and pd.notna(row[col]):
                value_2024 = row[col]
        
        if value_2021 is not None and value_2024 is not None and value_2021 > 0:
            pct_change = ((value_2024 - value_2021) / value_2021) * 100
            return pct_change, 2021, 2024, value_2021, value_2024
        
        return np.nan, np.nan, np.nan, np.nan, np.nan
    
    # Calculate percentage changes
    df1[['overall_pct_change', 'first_year', 'last_year', 'first_value', 'last_value']] = df1.apply(
        lambda row: pd.Series(calculate_overall_change(row)), axis=1
    )
    
    # Merge with shapefile
    gdf = shapefile.merge(df1, on=['FIRST_DNAM', 'FIRST_CHIE'], how='left', validate='1:1')
    
    # Fixed bins for percentage change (-70% to >+20%, 10 bins)
    def create_fixed_bins():
        # Fixed range from -70% to +20% with 10 equal bins
        bins = np.linspace(-70, 20, 11)  # 11 values create 10 bins
        return bins
    
    def create_map_with_legend(gdf_plot, title, filename, stats_text=None, show_dnam_boundaries=False, 
                              filter_to_data=False, show_names=False, name_column=None, use_simple_colors=False):
        """Helper function to create standardized maps"""
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        
        # Determine what to show as base
        if filter_to_data and len(gdf_plot) > 0:
            # Only show areas of interest - filter base map to regions with data
            if show_dnam_boundaries:
                # Get unique FIRST_DNAM values from the data
                dnam_values = gdf_plot['FIRST_DNAM'].dropna().unique()
                base_gdf = gdf[gdf['FIRST_DNAM'].isin(dnam_values)]
            else:
                # Just show the areas with data
                base_gdf = gdf_plot
            
            # Plot filtered base in light gray
            base_gdf.boundary.plot(ax=ax, color='lightgray', linewidth=0.5, alpha=0.3)
        else:
            # Plot full base shapefile in light gray
            gdf.boundary.plot(ax=ax, color='lightgray', linewidth=0.5, alpha=0.3)
        
        if len(gdf_plot) > 0:
            # Get valid data for binning
            valid_data = gdf_plot.dropna(subset=['overall_pct_change'])
            
            if len(valid_data) > 0:
                if use_simple_colors:
                    # Simple two-color scheme: #47B5FF for negative, pink for positive
                    colors = ['#47B5FF' if x < 0 else 'pink' for x in valid_data['overall_pct_change']]
                    
                    # Plot with simple colors
                    valid_data.plot(
                        color=colors,
                        ax=ax,
                        legend=False,
                        edgecolor='gray',
                        linewidth=1.0
                    )
                    
                    # Create simple legend
                    legend_elements = [
                        mpatches.Patch(color='#47B5FF', label='Decreasing Incidence'),
                        mpatches.Patch(color='pink', label='Increasing Incidence')
                    ]
                    
                    # Add legend
                    legend = ax.legend(handles=legend_elements, loc='center left', bbox_to_anchor=(1, 0.5), 
                                      title='Trend Direction', 
                                      title_fontsize=12, fontsize=10)
                    legend.get_title().set_fontweight('bold')
                else:
                    # Use fixed bins from -70% to +20%
                    bins = create_fixed_bins()
                    
                    # Create labels for fixed bins with >+20% for the last bin
                    bin_labels = []
                    for i in range(len(bins) - 1):
                        if i == len(bins) - 2:  # Last bin
                            bin_labels.append(f"{bins[i]:+.0f}% to >+{bins[i+1]:.0f}%")
                        else:
                            bin_labels.append(f"{bins[i]:+.0f}% to {bins[i+1]:+.0f}%")
                    
                    # Use RdBu_r: Red for increases, Blue for decreases
                    cmap = RdBu_r
                    norm = BoundaryNorm(bins, cmap.N)
                    
                    # Plot the data
                    valid_data.plot(
                        column='overall_pct_change',
                        ax=ax,
                        cmap=cmap,
                        norm=norm,
                        legend=False,
                        edgecolor='gray',
                        linewidth=1.0
                    )
                    
                    # Create custom legend with fixed bins
                    legend_elements = []
                    for i, (bin_label, color_val) in enumerate(zip(bin_labels, np.linspace(0, 1, len(bin_labels)))):
                        color = cmap(color_val)
                        legend_elements.append(mpatches.Patch(color=color, label=bin_label))
                    
                    # Add legend
                    legend = ax.legend(handles=legend_elements, loc='center left', bbox_to_anchor=(1, 0.5), 
                                      title='Overall % Change\n(Red=Increase, Blue=Decrease)', 
                                      title_fontsize=12, fontsize=10)
                    legend.get_title().set_fontweight('bold')
        
        # Add FIRST_DNAM boundaries if requested
        if show_dnam_boundaries:
            if filter_to_data and len(gdf_plot) > 0:
                # Only show boundaries for areas with data
                dnam_values = gdf_plot['FIRST_DNAM'].dropna().unique()
                dnam_boundaries = gdf[gdf['FIRST_DNAM'].isin(dnam_values)]
            else:
                dnam_boundaries = gdf
            
            # Group by FIRST_DNAM and plot boundaries
            for dnam in dnam_boundaries['FIRST_DNAM'].dropna().unique():
                dnam_geom = dnam_boundaries[dnam_boundaries['FIRST_DNAM'] == dnam]
                dnam_geom.boundary.plot(ax=ax, color='black', linewidth=2.5, alpha=1.0)
        
        # Add names if requested
        if show_names and name_column and len(gdf_plot) > 0:
            if name_column == 'FIRST_DNAM':
                # For FIRST_DNAM, show names only once at district center
                dnam_groups = gdf_plot.groupby('FIRST_DNAM')
                for dnam_name, group in dnam_groups:
                    # Calculate the centroid of all chiefdoms in this DNAM
                    union_geom = group.geometry.unary_union
                    if hasattr(union_geom, 'centroid'):
                        centroid = union_geom.centroid
                        ax.annotate(dnam_name, 
                                   xy=(centroid.x, centroid.y),
                                   xytext=(3, 3), textcoords="offset points",
                                   fontsize=10, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8),
                                   ha='center')
            elif name_column == 'FIRST_CHIE':
                # For FIRST_CHIE, show names with percentage change in brackets
                valid_data_names = gdf_plot.dropna(subset=[name_column, 'overall_pct_change'])
                
                for idx, row in valid_data_names.iterrows():
                    if hasattr(row.geometry, 'centroid'):
                        centroid = row.geometry.centroid
                        change_text = f"{row[name_column]} ({row['overall_pct_change']:+.1f}%)"
                        ax.annotate(change_text, 
                                   xy=(centroid.x, centroid.y),
                                   xytext=(3, 3), textcoords="offset points",
                                   fontsize=8, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.7),
                                   ha='left')
            else:
                # Default behavior for other name columns
                valid_data_names = gdf_plot.dropna(subset=[name_column])
                
                for idx, row in valid_data_names.iterrows():
                    if hasattr(row.geometry, 'centroid'):
                        centroid = row.geometry.centroid
                        ax.annotate(row[name_column], 
                                   xy=(centroid.x, centroid.y),
                                   xytext=(3, 3), textcoords="offset points",
                                   fontsize=8, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.7),
                                   ha='left')
        
        # Add title
        ax.set_title(title, fontsize=16, fontweight='bold', pad=20)
        ax.set_axis_off()
        
        # Add statistics if provided
        if stats_text:
            ax.text(0.02, 0.98, stats_text, transform=ax.transAxes, fontsize=11,
                    verticalalignment='top', bbox=dict(boxstyle='round', facecolor='white', alpha=0.8))
        
        # Save map
        map_path = os.path.join(output_dir, filename)
        plt.tight_layout()
        plt.savefig(map_path, dpi=400, bbox_inches='tight')
        plt.close()
        print(f"[Saved] {map_path}")
        return map_path
    
    # Remove rows with no percentage change data
    gdf_valid = gdf.dropna(subset=['overall_pct_change']).copy()
    
    if len(gdf_valid) == 0:
        print("Error: No valid percentage change data found")
        return
    
    # 1. NATIONAL OVERVIEW MAP with DNAM boundaries and names
    overall_avg = gdf_valid['overall_pct_change'].mean()
    overall_min = gdf_valid['overall_pct_change'].min()
    overall_max = gdf_valid['overall_pct_change'].max()
    total_chiefdoms = len(gdf_valid)
    first_year = int(gdf_valid['first_year'].mode().iloc[0])
    last_year = int(gdf_valid['last_year'].mode().iloc[0])
    
    national_stats = (f"Total Chiefdoms: {total_chiefdoms}\n"
                     f"National Average: {overall_avg:+.1f}%\n"
                     f"Range: {overall_min:+.1f}% to {overall_max:+.1f}%")
    
    create_map_with_legend(
        gdf_valid, 
        f"National Crude Incidence Change by Chiefdom (2021 to 2024)", 
        'national_crude_incidence_change_map.png',
        national_stats,
        show_dnam_boundaries=True,
        filter_to_data=True,
        show_names=True,
        name_column='FIRST_DNAM',
        use_simple_colors=True
    )
    
    # 2. INDIVIDUAL MAPS FOR EACH FIRST_DNAM (filtered to show only areas of interest)
    first_dnam_values = gdf_valid['FIRST_DNAM'].dropna().unique()
    
    for first_dnam in first_dnam_values:
        gdf_dnam = gdf_valid[gdf_valid['FIRST_DNAM'] == first_dnam].copy()
        
        if len(gdf_dnam) == 0:
            continue
        
        # Calculate statistics for this DNAM
        avg_change = gdf_dnam['overall_pct_change'].mean()
        min_change_dnam = gdf_dnam['overall_pct_change'].min()
        max_change_dnam = gdf_dnam['overall_pct_change'].max()
        n_chiefdoms = len(gdf_dnam)
        
        dnam_stats = (f"Chiefdoms: {n_chiefdoms}\n"
                     f"Average: {avg_change:+.1f}%\n"
                     f"Range: {min_change_dnam:+.1f}% to {max_change_dnam:+.1f}%")
        
        safe_dnam_name = "".join(c for c in first_dnam if c.isalnum() or c in (' ', '-', '_')).rstrip()
        
        create_map_with_legend(
            gdf_dnam,
            f"Crude Incidence Change - {first_dnam} (2021 to 2024)",
            f'crude_incidence_change_map_{safe_dnam_name}.png',
            dnam_stats,
            filter_to_data=True,
            show_names=True,
            name_column='FIRST_CHIE',
            use_simple_colors=True
        )
    
    # 3. COMBINED OVERVIEW MAP BY FIRST_DNAM (dissolved geometry)
    # Create dissolved dataframe for district-level mapping
    gdf_dissolved = gdf_valid.copy()
    # Calculate change per DNAM (2021 to 2024)
    dnam_change = gdf_dissolved.groupby('FIRST_DNAM')['overall_pct_change'].agg(['mean', 'count']).reset_index()
    dnam_change.columns = ['FIRST_DNAM', 'change_2021_2024', 'chiefdom_count']
    
    # Dissolve geometries by FIRST_DNAM to create district boundaries
    gdf_dnam_dissolved = gdf_dissolved.dissolve(by='FIRST_DNAM', aggfunc='first').reset_index()
    
    # Merge with change data
    gdf_dnam_dissolved = gdf_dnam_dissolved.merge(dnam_change, on='FIRST_DNAM')
    gdf_dnam_dissolved['overall_pct_change'] = gdf_dnam_dissolved['change_2021_2024']
    
    overview_stats = (f"Districts: {len(dnam_change)}\n"
                     f"Mean District Change: {dnam_change['change_2021_2024'].mean():+.1f}%\n"
                     f"Best District: {dnam_change['change_2021_2024'].max():+.1f}%\n"
                     f"Worst District: {dnam_change['change_2021_2024'].min():+.1f}%")
    
    create_map_with_legend(
        gdf_dnam_dissolved,
        f"Crude Incidence Change by District (2021 to 2024)",
        'district_change_overview_map.png',
        overview_stats,
        filter_to_data=True,
        show_names=True,
        name_column='FIRST_DNAM',
        use_simple_colors=True
    )
    
    # 4. MAPS FOR EACH FIRST_DNAM SHOWING FIRST_CHIE (filtered and enhanced)
    for first_dnam in first_dnam_values:
        dnam_df = df1[df1['FIRST_DNAM'] == first_dnam]
        
        if len(dnam_df) == 0:
            continue
        
        # Get unique FIRST_CHIE values for this FIRST_DNAM
        first_chie_values = dnam_df['FIRST_CHIE'].dropna().unique()
        
        if len(first_chie_values) == 0:
            continue
        
        # Filter geodataframe for this DNAM
        gdf_dnam_chie = gdf_valid[gdf_valid['FIRST_DNAM'] == first_dnam].copy()
        
        if len(gdf_dnam_chie) == 0:
            continue
        
        # Calculate statistics for CHIE within this DNAM
        chie_stats = gdf_dnam_chie.groupby('FIRST_CHIE')['overall_pct_change'].agg(['mean', 'count']).reset_index()
        chie_stats.columns = ['FIRST_CHIE', 'avg_change', 'count']
        
        stats_summary = (f"Chiefdoms in {first_dnam}: {len(gdf_dnam_chie)}\n"
                        f"Administrative Units: {len(chie_stats)}\n"
                        f"Avg Change: {gdf_dnam_chie['overall_pct_change'].mean():+.1f}%\n"
                        f"Range: {gdf_dnam_chie['overall_pct_change'].min():+.1f}% to {gdf_dnam_chie['overall_pct_change'].max():+.1f}%")
        
        safe_dnam_name = "".join(c for c in first_dnam if c.isalnum() or c in (' ', '-', '_')).rstrip()
        
        create_map_with_legend(
            gdf_dnam_chie,
            f"Chiefdom-Level Change within {first_dnam} (2021 to 2024)",
            f'chiefdom_level_change_map_{safe_dnam_name}.png',
            stats_summary,
            filter_to_data=True,
            show_names=True,
            name_column='FIRST_CHIE',
            use_simple_colors=True
        )
    
    print(f"\n=== COMPREHENSIVE TREND MAPPING COMPLETE ===")
    print(f"Created maps equivalent to all trend analyses:")
    print(f"1. National overview map with DNAM boundaries and names (filtered view)")
    print(f"2. {len(first_dnam_values)} individual DNAM maps (filtered to areas of interest)")
    print(f"3. District average overview map with names (filtered view)")
    print(f"4. {len(first_dnam_values)} DNAM chiefdom-level maps with CHIE names (filtered views)")
    print(f"\nColor scheme: RED = Increasing incidence, BLUE = Decreasing incidence")
    print(f"All maps saved in: {output_dir}")
    print(f"\nNational Summary:")
    print(f"- Total chiefdoms analyzed: {total_chiefdoms}")
    print(f"- National average change: {overall_avg:+.1f}%")
    print(f"- Time period: {first_year} to {last_year}")
    
    # Create summary statistics
    decreasing_count = len(gdf_valid[gdf_valid['overall_pct_change'] < 0])
    increasing_count = len(gdf_valid[gdf_valid['overall_pct_change'] > 0])
    stable_count = len(gdf_valid[gdf_valid['overall_pct_change'] == 0])
    
    print(f"\nTrend Distribution:")
    print(f"- Decreasing (blue): {decreasing_count} chiefdoms ({decreasing_count/total_chiefdoms*100:.1f}%)")
    print(f"- Increasing (red): {increasing_count} chiefdoms ({increasing_count/total_chiefdoms*100:.1f}%)")
    print(f"- Stable: {stable_count} chiefdoms ({stable_count/total_chiefdoms*100:.1f}%)")


def combine_xls(file_path):
    # Combine the files
    files = Path(file_path).glob("*.xls")
    df_list = [pd.read_excel(file) for file in files]
    combined_df = pd.concat(df_list, ignore_index=True)

    # Print head
    print("\n=== Preview of Combined Data ===")
    print(tabulate(combined_df.tail(), headers='keys', tablefmt='grid'))

    # Format column names into 3 columns
    print("\n=== Column Names (3 per row) ===")
    columns = list(combined_df.columns)
    padded_cols = columns + [""] * ((3 - len(columns) % 3) % 3)  # pad to multiple of 3
    col_table = [padded_cols[i:i+3] for i in range(0, len(padded_cols), 3)]
    print(tabulate(col_table, headers=["Column 1", "Column 2", "Column 3"], tablefmt="grid"))

    return combined_df

def combine_xlsx(file_path):
    # Combine the files
    files = Path(file_path).glob("*.xlsx")
    df_list = [pd.read_excel(file) for file in files]
    combined_df = pd.concat(df_list, ignore_index=True)

    # Print head
    print("\n=== Preview of Combined Data ===")
    print(tabulate(combined_df.tail(), headers='keys', tablefmt='grid'))

    # Format column names into 3 columns
    print("\n=== Column Names (3 per row) ===")
    columns = list(combined_df.columns)
    padded_cols = columns + [""] * ((3 - len(columns) % 3) % 3)  # pad to multiple of 3
    col_table = [padded_cols[i:i+3] for i in range(0, len(padded_cols), 3)]
    print(tabulate(col_table, headers=["Column 1", "Column 2", "Column 3"], tablefmt="grid"))

    return combined_df



def combine_csv(file_path):
    # Combine the files
    files = Path(file_path).glob("*.csv")
    df_list = [pd.read_csv(file) for file in files]
    combined_df = pd.concat(df_list, ignore_index=True)

    # Print head
    print("\n=== Preview of Combined Data ===")
    print(tabulate(combined_df.tail(), headers='keys', tablefmt='grid'))

    # Format column names into 3 columns
    print("\n=== Column Names (3 per row) ===")
    columns = list(combined_df.columns)
    padded_cols = columns + [""] * ((3 - len(columns) % 3) % 3)  # pad to multiple of 3
    col_table = [padded_cols[i:i+3] for i in range(0, len(padded_cols), 3)]
    print(tabulate(col_table, headers=["Column 1", "Column 2", "Column 3"], tablefmt="grid"))

    return combined_df



import pandas as pd
from pathlib import Path
from tabulate import tabulate
import warnings
from collections import defaultdict

import pandas as pd
from pathlib import Path
import warnings
from collections import defaultdict
import re
import unicodedata

def clean_column_name(col_name):
    """
    Clean column names by removing accents, special characters, and trailing spaces.
    
    Parameters:
    col_name (str): Original column name
    
    Returns:
    str: Cleaned column name
    """
    if pd.isna(col_name) or col_name == '':
        return col_name
    
    # Convert to string if not already
    col_name = str(col_name)
    
    # Remove trailing and leading spaces
    col_name = col_name.strip()
    
    # Remove accents (normalize to NFD, then remove combining characters)
    col_name = unicodedata.normalize('NFD', col_name)
    col_name = ''.join(char for char in col_name if unicodedata.category(char) != 'Mn')
    
    # Replace special characters with underscores (keep alphanumeric, spaces, and underscores)
    col_name = re.sub(r'[^\w\s]', '_', col_name)
    
    # Replace multiple spaces with single space
    col_name = re.sub(r'\s+', ' ', col_name)
    
    # Remove trailing and leading spaces again
    col_name = col_name.strip()
    
    return col_name

def clean_dataframe(df, filename):
    """
    Clean dataframe by removing unnamed columns and cleaning column names.
    
    Parameters:
    df (pd.DataFrame): Original dataframe
    filename (str): Name of the file for reporting
    
    Returns:
    pd.DataFrame: Cleaned dataframe
    """
    original_cols = len(df.columns)
    
    # Remove unnamed columns (columns that start with 'Unnamed:' or are empty/NaN)
    columns_to_keep = []
    unnamed_cols = []
    
    for col in df.columns:
        col_str = str(col)
        if (col_str.startswith('Unnamed:') or 
            pd.isna(col) or 
            col_str.strip() == '' or 
            col_str.strip() == 'nan'):
            unnamed_cols.append(col)
        else:
            columns_to_keep.append(col)
    
    # Keep only named columns
    if unnamed_cols:
        print(f"  Removed {len(unnamed_cols)} unnamed columns from {filename}")
        df = df[columns_to_keep]
    
    # Clean column names
    original_column_names = df.columns.tolist()
    cleaned_column_names = [clean_column_name(col) for col in df.columns]
    
    # Check if any column names were changed
    changed_cols = []
    for orig, clean in zip(original_column_names, cleaned_column_names):
        if str(orig) != clean:
            changed_cols.append((orig, clean))
    
    if changed_cols:
        print(f"  Cleaned {len(changed_cols)} column names in {filename}")
        for orig, clean in changed_cols[:3]:  # Show first 3 examples
            print(f"    '{orig}' -> '{clean}'")
        if len(changed_cols) > 3:
            print(f"    ... and {len(changed_cols) - 3} more")
    
    # Apply cleaned column names
    df.columns = cleaned_column_names
    
    # Remove duplicate column names by adding suffix
    if df.columns.duplicated().any():
        df.columns = pd.Index([f"{col}_{i}" if df.columns.tolist()[:i].count(col) > 0 else col 
                              for i, col in enumerate(df.columns)])
        print(f"  Resolved duplicate column names in {filename}")
    
    return df

import pandas as pd
from pathlib import Path
import warnings
from collections import defaultdict
import re
import unicodedata

def clean_column_name(col_name):
    """
    Clean column names by removing accents, special characters, and trailing spaces.
    
    Parameters:
    col_name (str): Original column name
    
    Returns:
    str: Cleaned column name
    """
    if pd.isna(col_name) or col_name == '':
        return col_name
    
    # Convert to string if not already
    col_name = str(col_name)
    
    # Remove trailing and leading spaces
    col_name = col_name.strip()
    
    # Remove accents (normalize to NFD, then remove combining characters)
    col_name = unicodedata.normalize('NFD', col_name)
    col_name = ''.join(char for char in col_name if unicodedata.category(char) != 'Mn')
    
    # Replace special characters with underscores (keep alphanumeric, spaces, and underscores)
    col_name = re.sub(r'[^\w\s]', '_', col_name)
    
    # Replace multiple spaces with single space
    col_name = re.sub(r'\s+', ' ', col_name)
    
    # Remove trailing and leading spaces again
    col_name = col_name.strip()
    
    return col_name

def clean_dataframe(df, filename):
    """
    Clean dataframe by removing unnamed columns and cleaning column names.
    
    Parameters:
    df (pd.DataFrame): Original dataframe
    filename (str): Name of the file for reporting
    
    Returns:
    pd.DataFrame: Cleaned dataframe
    """
    original_cols = len(df.columns)
    
    # Remove unnamed columns (columns that start with 'Unnamed:' or are empty/NaN)
    columns_to_keep = []
    unnamed_cols = []
    
    for col in df.columns:
        col_str = str(col)
        if (col_str.startswith('Unnamed:') or 
            pd.isna(col) or 
            col_str.strip() == '' or 
            col_str.strip() == 'nan'):
            unnamed_cols.append(col)
        else:
            columns_to_keep.append(col)
    
    # Keep only named columns
    if unnamed_cols:
        print(f"  Removed {len(unnamed_cols)} unnamed columns from {filename}")
        df = df[columns_to_keep]
    
    # Clean column names
    original_column_names = df.columns.tolist()
    cleaned_column_names = [clean_column_name(col) for col in df.columns]
    
    # Check if any column names were changed
    changed_cols = []
    for orig, clean in zip(original_column_names, cleaned_column_names):
        if str(orig) != clean:
            changed_cols.append((orig, clean))
    
    if changed_cols:
        print(f"  Cleaned {len(changed_cols)} column names in {filename}")
        for orig, clean in changed_cols[:3]:  # Show first 3 examples
            print(f"    '{orig}' -> '{clean}'")
        if len(changed_cols) > 3:
            print(f"    ... and {len(changed_cols) - 3} more")
    
    # Apply cleaned column names
    df.columns = cleaned_column_names
    
    # Remove duplicate column names by adding suffix
    if df.columns.duplicated().any():
        df.columns = pd.Index([f"{col}_{i}" if df.columns.tolist()[:i].count(col) > 0 else col 
                              for i, col in enumerate(df.columns)])
        print(f"  Resolved duplicate column names in {filename}")
    
    return df

def combine_files(file_path):
    """
    Combine files of the same type in a directory.
    Handles Excel (.xlsx, .xls), CSV (.csv), Stata (.dta), and SPSS (.sav) files.
    Cleans data by removing unnamed columns, trailing spaces, accents, and special characters.
    Manages both common and uncommon columns across files.
    
    Parameters:
    file_path (str): Path to directory containing files to combine
    
    Returns:
    pd.DataFrame: Combined dataframe with all files
    """
    
    file_path = Path(file_path)
    
    # Define file readers for different formats
    readers = {
        '.xlsx': pd.read_excel,
        '.xls': pd.read_excel,
        '.csv': pd.read_csv,
        '.dta': pd.read_stata,
        '.sav': pd.read_spss
    }
    
    # Find all files and determine the file type
    all_files = list(file_path.glob("*"))
    data_files = []
    file_extension = None
    
    # Get the first data file extension to determine file type
    for file in all_files:
        if file.suffix.lower() in readers:
            if file_extension is None:
                file_extension = file.suffix.lower()
            elif file.suffix.lower() != file_extension:
                print(f"Warning: Mixed file types found. Using {file_extension} files only.")
                continue
            data_files.append(file)
    
    if not data_files:
        raise ValueError("No supported data files found in the directory")
    
    if file_extension not in readers:
        raise ValueError(f"Unsupported file format: {file_extension}")
    
    print(f"Found {len(data_files)} {file_extension} files to combine")
    
    # Read all files and track columns
    df_list = []
    file_columns = {}  # Track columns for each file
    all_columns = set()
    
    reader_func = readers[file_extension]
    
    for file in data_files:
        try:
            print(f"Reading: {file.name}")
            
            # Special handling for different file types
            if file_extension == '.csv':
                df = reader_func(file, encoding='utf-8')
            elif file_extension == '.sav':
                df = reader_func(file, apply_value_formats=True)
            else:
                df = reader_func(file)
            
            # Clean the dataframe
            df = clean_dataframe(df, file.name)
            
            # Skip if dataframe is empty after cleaning
            if df.empty:
                print(f"  Skipping {file.name}: No data remaining after cleaning")
                continue
            
            df_list.append(df)
            file_columns[file.name] = set(df.columns)
            all_columns.update(df.columns)
            
        except Exception as e:
            print(f"Error reading {file.name}: {str(e)}")
            continue
    
    if not df_list:
        raise ValueError("No files could be read successfully")
    
    # Analyze column differences
    print("\n=== Data Cleaning Summary ===")
    print("All files have been cleaned:")
    print("- Removed unnamed/empty columns")
    print("- Cleaned column names (removed accents, special characters, trailing spaces)")
    print("- Resolved duplicate column names")
    
    print("\n=== Column Analysis ===")
    
    # Find common columns (present in all files)
    common_columns = set(file_columns[list(file_columns.keys())[0]])
    for file_cols in file_columns.values():
        common_columns &= file_cols
    
    print(f"Common columns across all files: {len(common_columns)}")
    if common_columns:
        print("Common columns:", sorted(list(common_columns)))
    
    # Find uncommon columns and report which files are missing them
    uncommon_columns = all_columns - common_columns
    if uncommon_columns:
        print(f"\nUncommon columns found: {len(uncommon_columns)}")
        
        # Create a detailed report of missing columns
        column_file_map = defaultdict(list)
        for col in uncommon_columns:
            files_with_col = []
            files_without_col = []
            
            for filename, cols in file_columns.items():
                if col in cols:
                    files_with_col.append(filename)
                else:
                    files_without_col.append(filename)
            
            print(f"\nColumn '{col}':")
            print(f"  Present in: {', '.join(files_with_col)}")
            if files_without_col:
                print(f"  Missing from: {', '.join(files_without_col)}")
    
    # Combine all dataframes (pandas will handle missing columns by filling with NaN)
    print(f"\n=== Combining Files ===")
    combined_df = pd.concat(df_list, ignore_index=True, sort=False)
    
    print(f"Combined dataset shape: {combined_df.shape}")
    print(f"Total columns in combined dataset: {len(combined_df.columns)}")
    
    # Summary statistics
    print(f"\n=== Summary ===")
    print(f"Files combined: {len(df_list)}")
    print(f"Total rows: {len(combined_df)}")
    print(f"Total columns: {len(combined_df.columns)}")
    print(f"Common columns: {len(common_columns)}")
    print(f"Uncommon columns: {len(uncommon_columns)}")
    
    # Check for missing values in uncommon columns
    if uncommon_columns:
        missing_summary = []
        for col in uncommon_columns:
            missing_count = combined_df[col].isna().sum()
            missing_pct = (missing_count / len(combined_df)) * 100
            missing_summary.append([col, missing_count, f"{missing_pct:.1f}%"])
        
        print("\n=== Missing Values in Uncommon Columns ===")
        for col, missing_count, missing_pct in missing_summary:
            print(f"{col:<30} Missing: {missing_count:<8} ({missing_pct})")
    
    # Save the combined data to Excel
    output_filename = file_path / "combined_data.xlsx"
    try:
        combined_df.to_excel(output_filename, index=False)
        print(f"\n=== File Saved ===")
        print(f"Combined data saved to: {output_filename}")
        print(f"File is ready for download!")
    except Exception as e:
        print(f"\nError saving file: {str(e)}")
        print("Trying alternative filename...")
        try:
            alt_filename = file_path / f"combined_data_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            combined_df.to_excel(alt_filename, index=False)
            print(f"Combined data saved to: {alt_filename}")
        except Exception as e2:
            print(f"Failed to save file: {str(e2)}")
    
    return combined_df

def rename(df):
    name_map = pd.read_excel("input_files/others/old_new_rename.xlsx")
    rename_dict = dict(zip(name_map.iloc[:, 0], name_map.iloc[:, 1]))
    return df.rename(columns=rename_dict)
    
def compute(df):
    try:
        comp = pd.read_excel("input_files/others/compute new variables_python.xlsx")
    except Exception as e:
        raise FileNotFoundError(f"Error reading compute file: {e}")

    for i in range(len(comp)):
        new_var = comp.at[i, 'new_variable']
        op = comp.at[i, 'operation']
        components = [x.strip() for x in comp.at[i, 'components'].split(',')]

        # Check if all required columns exist in df
        missing_cols = [col for col in components if col not in df.columns]
        if missing_cols:
            print(f"Skipping '{new_var}' — missing columns: {missing_cols}")
            continue

        if op == "add":
            df[new_var] = df[components].sum(axis=1, skipna=True, min_count=1)
        elif op == "subtract" and len(components) >= 2:
            df[new_var] = df[components[0]] - df[components[1]]
            df[new_var] = df[new_var].clip(lower=0)
        else:
            print(f"Skipping '{new_var}' — unsupported operation or insufficient components.")
    
    return df



def sort(df):
    try:
        comp = pd.read_excel("input_files/others/compute new variables_python.xlsx")
    except Exception as e:
        raise FileNotFoundError(f"Could not read compute file: {e}")

    sorted_columns = []

    # Collect components and new variables in order
    for i in range(len(comp)):
        components = [x.strip() for x in str(comp.at[i, 'components']).split(',')]
        new_var = comp.at[i, 'new_variable']
        sorted_columns.extend(components)
        sorted_columns.append(new_var)

    # Ensure uniqueness and keep only existing columns
    sorted_columns = [col for col in dict.fromkeys(sorted_columns) if col in df.columns]

    # Add any remaining columns not in the sort list
    remaining_columns = [col for col in df.columns if col not in sorted_columns]

    # Final column order
    final_order = remaining_columns + sorted_columns

    # Reorder DataFrame
    return df[final_order]


def split(df):
    try:
        # Read the mapping file
        mapping = pd.read_excel("input_files/others/split colums.xlsx")
    except Exception as e:
        raise FileNotFoundError(f"Could not read split columns file: {e}")

    # Validate expected columns
    required_cols = {'original_col', 'new_col_month', 'new_col_year'}
    if not required_cols.issubset(mapping.columns):
        raise ValueError(f"Missing expected columns in mapping file: {required_cols - set(mapping.columns)}")

    original_col = mapping.at[0, 'original_col']
    new_col_month = mapping.at[0, 'new_col_month']
    new_col_year = mapping.at[0, 'new_col_year']

    # Ensure original column exists in df
    if original_col not in df.columns:
        raise KeyError(f"Column '{original_col}' not found in DataFrame.")

    # Split the column
    split_data = df[original_col].astype(str).str.strip().str.split(' ', n=1, expand=True)

    # Assign new columns
    df[new_col_month] = split_data[0].str.strip()
    df[new_col_year] = split_data[1].str.strip() if split_data.shape[1] > 1 else None

    # Month name to number mapping (English and French)
    month_map = {
        'January': '01', 'Janvier': '01',
        'February': '02', 'Février': '02', 'Fevrier': '02',
        'March': '03', 'Mars': '03',
        'April': '04', 'Avril': '04',
        'May': '05', 'Mai': '05',
        'June': '06', 'Juin': '06',
        'July': '07', 'Juillet': '07',
        'August': '08', 'Août': '08', 'Aout': '08',
        'September': '09', 'Septembre': '09',
        'October': '10', 'Octobre': '10',
        'November': '11', 'Novembre': '11',
        'December': '12', 'Décembre': '12', 'Decembre': '12'
    }

    # Standardize month
    df[new_col_month] = df[new_col_month].map(lambda x: month_map.get(x, x))

    return df


### Outlier detection and correctio with winsorized method
import pandas as pd
import numpy as np

# Function to detect outliers using Scatterplot with Q1 and Q3 lines
def detect_outliers_scatterplot(df, col):
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return lower_bound, upper_bound

# Function to apply winsorization to a column
def winsorize_series(series, lower_bound, upper_bound):
    return series.clip(lower=lower_bound, upper=upper_bound)

# Function to process a single column (grouped by adm1, adm2, adm3, hf, year)
def process_column_winsorization(df, column):
    grouped = df.groupby(['adm1', 'adm2', 'adm3', 'hf', 'year'])
    results = []

    for (adm1, adm2, adm3, hf, year), group in grouped:
        group = group.copy()
        lower_bound, upper_bound = detect_outliers_scatterplot(group, column)

        group[f'{column}_lower_bound'] = lower_bound
        group[f'{column}_upper_bound'] = upper_bound
        group[f'{column}_category'] = np.where(
            (group[column] < lower_bound) | (group[column] > upper_bound),
            'Outlier',
            'Non-Outlier'
        )
        group[f'{column}_winsorized'] = winsorize_series(group[column], lower_bound, upper_bound)

        results.append(group)

    final_df = pd.concat(results)

    export_columns = [
        'adm1', 'adm2', 'adm3', 'hf', 'year', 'month', column,
        f'{column}_category', f'{column}_lower_bound', f'{column}_upper_bound',
        f'{column}_winsorized'
    ]
    export_columns = [col for col in export_columns if col in final_df.columns]

    return final_df[export_columns]

# Main function to process multiple columns and merge the results
def detect_outliers(df):
    columns_to_process = ['allout', 'susp', 'test', 'conf', 'maltreat', 'pres', 'maladm', 'maldth']
    processed_dfs = []

    for column in columns_to_process:
        if column not in df.columns:
            continue
        if df[column].isnull().all():
            continue

        processed_df = process_column_winsorization(df, column)
        processed_dfs.append(processed_df)

    if processed_dfs:
        merge_keys = ['adm1', 'adm2', 'adm3', 'hf', 'year', 'month']
        final_combined_df = processed_dfs[0]

        for df_to_merge in processed_dfs[1:]:
            final_combined_df = final_combined_df.merge(
                df_to_merge, on=merge_keys, how='outer', suffixes=('', '_dup')
            )
            final_combined_df = final_combined_df[[col for col in final_combined_df.columns if not col.endswith('_dup')]]

        return final_combined_df
    else:
        return None

import pandas as pd
from tabulate import tabulate

def outlier_summary(df):
    # Automatically detect columns ending with '_category'
    category_columns = [col for col in df.columns if col.endswith('_category')]
    
    summary_stats = {}

    for col in category_columns:
        total_outliers = (df[col] == 'Outlier').sum()
        total_non_outliers = (df[col] == 'Non-Outlier').sum()
        total = total_outliers + total_non_outliers

        if total > 0:
            outlier_percentage = (total_outliers / total) * 100
            non_outlier_percentage = (total_non_outliers / total) * 100
        else:
            outlier_percentage = 0
            non_outlier_percentage = 0

        summary_stats[col] = {
            'Total Outliers': total_outliers,
            'Total Non-Outliers': total_non_outliers,
            'Total Records': total,
            'Outlier Percentage': f"{outlier_percentage:.2f}%",
            'Non-Outlier Percentage': f"{non_outlier_percentage:.2f}%"
        }

    summary_df = pd.DataFrame(summary_stats).T

    # Print in a pretty table
    print(tabulate(summary_df, headers='keys', tablefmt='pretty'))

    return summary_df

### Outlier detection after correction with winsorized method
import pandas as pd
import numpy as np

# Function to detect outliers using Scatterplot with Q1 and Q3 lines
def detect_outliers_scatterplot(df, col):
    
    # Calculate Q1 and Q3
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    
    # Calculate the lower and upper bounds for outliers
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    
    return lower_bound, upper_bound

# Function to apply winsorization to a column
def winsorize_series(series, lower_bound, upper_bound):
  
    # Clip the values that are outside the bounds
    return series.clip(lower=lower_bound, upper=upper_bound)

# Function to process a column and return a DataFrame with winsorized data
def process_column_winsorization(df, column):
 
    # Group by 'adm1', 'adm2', 'adm3', 'hf', 'year' for processing each group separately
    grouped = df.groupby(['adm1', 'adm2', 'adm3', 'hf', 'year'])
    results = []

    # Process each group
    for (adm1, adm2, adm3, hf, year), group in grouped:
        # Detect outliers
        lower_bound, upper_bound = detect_outliers_scatterplot(group, column)
        
        # Add new columns for outlier boundaries, category, and winsorized data
        group[f'{column}_lower_bound'] = lower_bound
        group[f'{column}_upper_bound'] = upper_bound
        group[f'{column}_category'] = np.where(
            (group[column] < lower_bound) | (group[column] > upper_bound), 'Outlier', 'Non-Outlier'
        )

        
        # Append the processed group to the results list
        results.append(group)

    # Concatenate all the processed groups
    final_df = pd.concat(results)
    
    # Define the columns to export
    export_columns = [
        'adm1', 'adm2', 'adm3', 'hf', 'year', 'month', column,
        f'{column}_category', f'{column}_lower_bound', f'{column}_upper_bound',
       
    ]
    
    # Filter to include only the existing columns in the DataFrame
    export_columns = [col for col in export_columns if col in final_df.columns]
    
    return final_df[export_columns]

# Main function to process multiple columns and merge the results
def detect_outliers_after_correction(df):
    # List of columns to process
    columns_to_process = ['allout_winsorized', 'susp_winsorized', 'test_winsorized', 'conf_winsorized', 'maltreat_winsorized', 'pres_winsorized', 'maladm_winsorized', 'maldth_winsorized']
    processed_dfs = []

    # Loop through each column and process it
    for column in columns_to_process:
        if column not in df.columns:
            print(f"Skipping column {column} as it does not exist in the dataset.")
            continue
        if df[column].isnull().all():
            print(f"Skipping column {column} as it contains only missing values.")
            continue

        print(f"Processing column: {column}")
        processed_df = process_column_winsorization(df, column)
        processed_dfs.append(processed_df)

    # Merge the processed DataFrames
    if processed_dfs:
        merge_keys = ['adm1', 'adm2', 'adm3', 'hf', 'year', 'month']
        final_combined_df = processed_dfs[0]
        for df_to_merge in processed_dfs[1:]:
            final_combined_df = final_combined_df.merge(df_to_merge, on=merge_keys, how='outer')
        
        return final_combined_df
    else:
        print("No valid columns were processed.")
        return None

import pandas as pd
from tabulate import tabulate

def outlier_summary_after_correction(df):
    # Automatically detect columns ending with '_category'
    category_columns = [col for col in df.columns if col.endswith('_category')]
    
    summary_stats = {}

    for col in category_columns:
        total_outliers = (df[col] == 'Outlier').sum()
        total_non_outliers = (df[col] == 'Non-Outlier').sum()
        total = total_outliers + total_non_outliers

        if total > 0:
            outlier_percentage = (total_outliers / total) * 100
            non_outlier_percentage = (total_non_outliers / total) * 100
        else:
            outlier_percentage = 0
            non_outlier_percentage = 0

        summary_stats[col] = {
            'Total Outliers': total_outliers,
            'Total Non-Outliers': total_non_outliers,
            'Total Records': total,
            'Outlier Percentage': f"{outlier_percentage:.2f}%",
            'Non-Outlier Percentage': f"{non_outlier_percentage:.2f}%"
        }

    summary_df = pd.DataFrame(summary_stats).T

    # Print in a pretty table
    print(tabulate(summary_df, headers='keys', tablefmt='pretty'))

    return summary_df


# Epi stratification

import pandas as pd
import numpy as np
import geopandas as gpd
from functools import reduce
import os

def epi_stratification(
    output_folder='epi_output',
    output_filename='adjusted_incidence_with_mean_median.xlsx'
):
    # Create output directory if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    output_file = os.path.join(output_folder, output_filename)

    # Load input data
    routine_data = pd.read_excel("input_files/routine/clean_data/routine_data (1).xlsx")
    population_data = pd.read_excel("input_files/routine/population_data/population_data.xlsx")
    df = routine_data.copy()

    # Preprocess dates
    df['date'] = df['year'].astype(str) + '-' + df['month'].astype(str).str.zfill(2)
    df['date'] = pd.to_datetime(df['date'], format='%Y-%m').dt.to_period('M')
    df['Total_Reports'] = df[['allout', 'susp', 'test', 'conf', 'maltreat']].sum(axis=1)

    # Dynamic year range
    start = df['date'].min().year
    end = df['date'].max().year
    years = range(start, end + 1)

    # First reporting date
    df_active = df[df['Total_Reports'] > 0]
    first_report_dates = df_active.groupby(['adm1', 'adm2', 'adm3', 'hf'])['date'].min().reset_index()
    first_report_dates.rename(columns={'date': 'First_Reported_Date'}, inplace=True)

    # Reporting stats
    reporting_stats_by_year = []
    for year in years:
        df_year = df[df['year'] == year]
        reported = (
            df_year[df_year['conf'] > 0]
            .groupby(['adm1', 'adm2', 'adm3'], as_index=False)['conf']
            .count()
            .rename(columns={'conf': f'Times_Reported_{year}'})
        )
        expected = (
            first_report_dates
            .assign(Times_Expected=lambda x: np.where(
                x['First_Reported_Date'].dt.year == year,
                12 - x['First_Reported_Date'].dt.month + 1,
                np.where(year > x['First_Reported_Date'].dt.year, 12, 0)
            ))
            .groupby(['adm1', 'adm2', 'adm3'], as_index=False)['Times_Expected']
            .sum()
            .rename(columns={'Times_Expected': f'Times_Expected_To_Report_{year}'})
        )
        stats = pd.merge(expected, reported, on=['adm1', 'adm2', 'adm3'], how='outer')
        stats[f'Times_Reported_{year}'] = stats[f'Times_Reported_{year}'].fillna(0)
        stats[f'Times_Expected_To_Report_{year}'] = stats[f'Times_Expected_To_Report_{year}'].fillna(0)
        stats[f'conf_RR_{year}'] = (
            stats[f'Times_Reported_{year}']
            .div(stats[f'Times_Expected_To_Report_{year}'])
            .replace([np.inf, -np.inf], 0)
            .fillna(0)
            .round(2)
        )
        reporting_stats_by_year.append(stats)

    confirmed_data = reduce(
        lambda left, right: pd.merge(left, right, on=['adm1', 'adm2', 'adm3'], how='outer'),
        reporting_stats_by_year
    )

    # Aggregated routine data by year
    dfs = []
    for year in years:
        df_year = df[df['year'] == year]
        grouped = df_year.groupby(['adm1', 'adm2', 'adm3'], as_index=False)[['conf', 'test', 'pres']].sum()
        grouped = grouped.rename(columns={
            'conf': f'conf_{year}', 'test': f'test_{year}', 'pres': f'pres_{year}'
        })
        dfs.append(grouped)

    df_merge = reduce(lambda left, right: pd.merge(left, right, on=['adm1', 'adm2', 'adm3'], how='outer'), dfs)

    # Merge all
    df1 = df_merge.merge(confirmed_data, on=['adm1', 'adm2', 'adm3'], how='left', validate='1:1')
    data = df1.merge(population_data, on='adm3', how='left', validate='1:1')
  
    # Compute metrics
    for year in years:
        conf_col = f"conf_{year}"
        test_col = f"test_{year}"
        pop_col = f"pop{year}"
        pres_col = f"pres_{year}"
        conf_RR_col = f"conf_RR_{year}"

        if not all(col in data.columns for col in [conf_col, test_col, pop_col, pres_col, conf_RR_col]):
            continue

        # Test positivity rate (as proportion)
        data[f'TPR_{year}'] = data[conf_col].div(data[test_col])
    
        # Crude incidence rate per 1000 population
        data[f'crude_incidence_{year}'] = data[conf_col].div(data[pop_col]) * 1000
    
        # Calculate N1 (adjusted for presumed cases)
        data[f'N1_{year}'] = data[conf_col] + (data[pres_col] * data[f'TPR_{year}'])
    
        # Calculate N2 (adjusted for reporting rate)
        data[f'N2_{year}'] = data[f'N1_{year}'].div(data[conf_RR_col])
    
        # Calculate private facility adjustment
        private_adjustment = (data[f'N2_{year}'] * data['CSpr']).div(data['CSpu'])
    
        # Calculate non-facility adjustment
        non_facility_adjustment = (data[f'N2_{year}'] * data['CSn']).div(data['CSpu'])
    
        # Calculate N3 (total adjusted cases)
        data[f'N3_{year}'] = data[f'N2_{year}'] + private_adjustment + non_facility_adjustment
    
        # Calculate adjusted incidence rates per 1000 population
        data[f'adjusted1_{year}'] = data[f'N1_{year}'].div(data[pop_col]) * 1000
        data[f'adjusted2_{year}'] = data[f'N2_{year}'].div(data[pop_col]) * 1000
        data[f'adjusted3_{year}'] = data[f'N3_{year}'].div(data[pop_col]) * 1000
        

    # Summary stats
    for prefix in ['adjusted1', 'adjusted2', 'adjusted3']:
        cols = [f'{prefix}_{year}' for year in years if f'{prefix}_{year}' in data.columns]
        data[f'{prefix}_mean'] = data[cols].mean(axis=1)
        data[f'{prefix}_median'] = data[cols].median(axis=1)

    # Save
    data.to_excel(output_file, index=False)
    print(f"Data has been successfully saved to {output_folder}")
    return data

# Epi plots (individual)
import os
import re
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from matplotlib.colors import BoundaryNorm
from matplotlib.patches import Patch
import numpy as np

def merge_data_with_shapefile(df1, shapefile):
    gdf = shapefile.merge(df1, on=['FIRST_DNAM', 'FIRST_CHIE'], how='left', validate='1:1')
    return gdf

def individual_plots(
                     prefixes=['crude_incidence_', 'adjusted1_', 'adjusted2_', 'adjusted3_'],
                     colormap='RdYlBu_r',
                     edge_color='gray',
                     bins=[0, 50, 100, 250, 450, 700, 1000, float('inf')],
                     bin_labels=['<50', '50-100', '100-250', '250-450', '450-700', '700-1000', '>1000'],
                     output_root='epi_maps'):
    """
    Creates individual maps for each column with a valid prefix and 4-digit year.
    Saves each map in a subfolder named after the prefix inside the 'epi_maps' folder.
    """

    # Load input data
    df1 = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    os.makedirs(output_root, exist_ok=True)

    # Merge data
    gdf = merge_data_with_shapefile(df1, shapefile)

    # Detect valid columns
    pattern = re.compile(r'_(\d{4})$')
    columns_to_plot = []
    for col in gdf.columns:
        for prefix in prefixes:
            if col.startswith(prefix) and pattern.search(col):
                columns_to_plot.append((col, prefix))
                break

    if not columns_to_plot:
        print("No valid columns found.")
        return

    # Setup color map
    cmap = plt.cm.get_cmap(colormap, len(bins) - 1)
    norm = BoundaryNorm(bins, ncolors=cmap.N)

    for column_name, prefix in columns_to_plot:
        fig, ax = plt.subplots(figsize=(10, 10))

        valid_data = gdf[column_name].dropna()
        counts, _ = np.histogram(valid_data, bins=bins)
        bin_labels_with_counts = [f"{label} ({count})" for label, count in zip(bin_labels, counts)]

        gdf.plot(
            column=column_name,
            cmap=cmap,
            norm=norm,
            edgecolor=edge_color,
            linewidth=0.5,
            legend=False,
            ax=ax,
            missing_kwds={'color': 'lightgrey', 'edgecolor': 'white', 'linewidth': 0.3}
        )

        gdf.dissolve(by="FIRST_DNAM").boundary.plot(ax=ax, color="black", linewidth=1)      
     

        legend_elements = [
            Patch(facecolor=cmap(norm(bin_start)), edgecolor='black', label=label)
            for bin_start, label in zip(bins[:-1], bin_labels_with_counts)
        ]

        ax.legend(
            handles=legend_elements,
            loc='lower right',
            title="Cases per 1000",
            fontsize=9,
            title_fontsize=10,
            frameon=True,
            framealpha=1.0,
            ncol=1
        )

        ax.set_title(column_name.replace("_", " "), fontsize=14, pad=10)
        ax.axis("off")

        # Subfolder for each prefix
        prefix_folder = os.path.join(output_root, prefix.rstrip("_"))
        os.makedirs(prefix_folder, exist_ok=True)

        output_file = os.path.join(prefix_folder, f"{column_name}.png")
        plt.savefig(output_file, dpi=300, bbox_inches="tight")
        print(f"Saved: {column_name}.png to {prefix_folder}")
        plt.close()


## Subplots
import os
import re
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from matplotlib.colors import BoundaryNorm
from matplotlib.patches import Patch
import numpy as np

def subplots():
    prefixes = ['crude_incidence_', 'adjusted1_', 'adjusted2_', 'adjusted3_']
    os.makedirs("subplots", exist_ok=True)
    
    df1 = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    gdf_shape = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    gdf = gdf_shape.merge(df1, on=["FIRST_DNAM", "FIRST_CHIE"], how="left", validate="1:1")
    
    bins = [0, 50, 100, 250, 450, 700, 1000, float("inf")]
    labels = ['<50', '50-100', '100-250', '250-450', '450-700', '700-1000', '>1000']
    cmap = plt.cm.get_cmap("RdYlBu_r", len(bins)-1)
    norm = BoundaryNorm(bins, cmap.N)

    for prefix in prefixes:
        pattern = re.compile(f"^{re.escape(prefix)}(\\d{{4}})$")
        columns = [(col, pattern.match(col).group(1)) for col in gdf.columns if pattern.match(col)]

        if not columns:
            print(f"[Skipped] No columns found for prefix '{prefix}'")
            continue

        columns.sort(key=lambda x: x[1])

        fig, axes = plt.subplots(2, 5, figsize=(24, 15))
        axes = axes.flatten()

        for i in range(len(columns), 9):
            axes[i].set_visible(False)

        for i, ((col, year), ax) in enumerate(zip(columns, axes)):
            gdf.plot(
                column=col,
                cmap=cmap,
                norm=norm,
                edgecolor='gray',
                linewidth=0.5,
                ax=ax,
                legend=False,
                missing_kwds={"color": "lightgrey"}
            )

            gdf.dissolve(by="FIRST_DNAM").boundary.plot(ax=ax, color="black", linewidth=1)
            ax.set_title(year, fontsize=11)
            ax.axis("off")

            # Create and add legend for each plot
            data = gdf[col].dropna()
            counts, _ = np.histogram(data, bins=bins)
            legend_labels = [f"{label} ({count})" for label, count in zip(labels, counts)]
            legend_items = [
                Patch(facecolor=cmap(norm(b)), edgecolor='black', label=lab)
                for b, lab in zip(bins[:-1], legend_labels)
            ]

            ax.legend(
                handles=legend_items,
                fontsize=7,
                title="Cases/1000",
                loc='upper center',
                bbox_to_anchor=(0.5, -0.15),
                frameon=True
            )

        plt.subplots_adjust(wspace=0.5, hspace=0.4, right=0.9)
        output_path = f"subplots/{prefix.rstrip('_')}_maps.png"
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        print(f"[Saved] {output_path}")

        
## Line plots
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import re
import os

def epi_trends(output_folder='epi_lineplots'):
    os.makedirs(output_folder, exist_ok=True)

    # Read the Excel file
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")

    # Define prefixes and colors
    prefixes = ['crude_incidence', 'adjusted1', 'adjusted2', 'adjusted3']
    colors = ['blue', 'green', 'orange', 'red']

    # Get list of years from column names
    pattern = re.compile(rf'^({"|".join(prefixes)})_(\d{{4}})$')
    # Extract years from column names that match the pattern
    years = sorted(set(int(match.group(2)) for col in df.columns if (match := pattern.match(col))))

    # Loop through each district (adm1 = FIRST_DNAM)
    for district in df['FIRST_DNAM'].dropna().unique():
        df_district = df[df['FIRST_DNAM'] == district]
        chiefdoms = df_district['FIRST_CHIE'].dropna().unique()
        n = len(chiefdoms)

        n_cols = 4
        n_rows = int(np.ceil(n / n_cols))

        fig, axes = plt.subplots(n_rows, n_cols, figsize=(n_cols * 5, n_rows * 4), sharex=True, sharey=True)
        axes = axes.flatten()

        for i, chiefdom in enumerate(chiefdoms):
            ax = axes[i]
            row = df_district[df_district['FIRST_CHIE'] == chiefdom]

            if row.empty:
                ax.set_title(f"{chiefdom} (No data)")
                ax.axis("off")
                continue

            for prefix, color in zip(prefixes, colors):
                cols = [f"{prefix}_{year}" for year in years if f"{prefix}_{year}" in row.columns]
                values = row[cols].values.flatten()

                if len(values) != len(years) or all(pd.isna(values)):
                    continue  # Skip if no valid data

                ax.plot(years, values, marker='o', label=prefix.replace('_', ' ').title(), color=color)

            ax.set_title(chiefdom, fontsize=10)
            ax.tick_params(axis='x', rotation=45)

        # Turn off unused axes
        for j in range(i + 1, len(axes)):
            axes[j].axis("off")

        # Shared legend
        handles, labels = axes[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, title="Indicator", loc="lower center", ncol=4)

        fig.suptitle(f"Incidence Trends by Chiefdom - {district}", fontsize=14)
        plt.tight_layout(rect=[0, 0.05, 1, 0.95])

        filename = os.path.join(output_folder, f"{district}_trends.png")
        plt.savefig(filename, dpi=300)
        plt.close()
        print(f"[Saved] {filename}")

    return df


## Crude incidence trends
import os
import re
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt


def crude_trends(output_folder='crude_plots'):
    os.makedirs(output_folder, exist_ok=True)

    # Read the Excel file
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")

    # Define prefixes and colors
    prefixes = ['crude_incidence']
    colors = ['blue']

    # Get list of years from column names using the first prefix
    pattern = re.compile(r'^crude_incidence_(\d{4})$')
    years = sorted(int(pattern.match(col).group(1)) for col in df.columns if pattern.match(col))

    # Loop through each district (adm1 = FIRST_DNAM)
    for district in df['FIRST_DNAM'].dropna().unique():
        df_district = df[df['FIRST_DNAM'] == district]
        chiefdoms = df_district['FIRST_CHIE'].dropna().unique()
        n = len(chiefdoms)

        n_cols = 4
        n_rows = int(np.ceil(n / n_cols))

        fig, axes = plt.subplots(n_rows, n_cols, figsize=(n_cols * 6, n_rows * 5), sharex=False, sharey=True)
        axes = axes.flatten()

        for i, chiefdom in enumerate(chiefdoms):
            ax = axes[i]
            row = df_district[df_district['FIRST_CHIE'] == chiefdom]

            if row.empty:
                ax.set_title(f"{chiefdom} (No data)", fontweight='bold')
                ax.axis("off")
                continue

            for prefix, color in zip(prefixes, colors):
                cols = [f"{prefix}_{year}" for year in years if f"{prefix}_{year}" in row.columns]
                values = row[cols].values.flatten()

                if len(values) != len(years) or all(pd.isna(values)):
                    continue  # Skip if no valid data

                ax.plot(years, values, marker='o', label=prefix.replace('_', ' ').title(), color=color, linewidth=2)

                # Add trend line if enough data
                if np.count_nonzero(~np.isnan(values)) >= 2:
                    fit = np.polyfit(years, values, 1)
                    trend_line = np.poly1d(fit)(years)
                    ax.plot(years, trend_line, linestyle='--', color=color, alpha=0.7, label="Trend", linewidth=2)

            ax.set_title(chiefdom, fontsize=11, fontweight='bold')
            ax.set_xlabel("Year", fontsize=10, fontweight='bold')
            ax.set_ylabel("Incidence", fontsize=10, fontweight='bold')
            ax.grid(True)
            ax.tick_params(axis='x', rotation=0)
            ax.tick_params(axis='both', labelsize=9)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontweight('bold')

        # Turn off unused subplots
        for j in range(i + 1, len(axes)):
            axes[j].axis("off")

        # Add shared legend
        handles, labels = axes[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, title="Indicator", loc="lower center", ncol=4, fontsize=10, title_fontsize=11, prop={'weight': 'bold'})

        fig.suptitle(f"Crude Incidence Trends by Chiefdom - {district}", fontsize=16, fontweight='bold')
        plt.tight_layout(rect=[0, 0.05, 1, 0.95])

        filename = os.path.join(output_folder, f"{district}_trends.png")
        plt.savefig(filename, dpi=400, bbox_inches='tight')
        plt.close()
        print(f"[Saved] {filename}")

    return df

## Adjusted1 trend

import os
import re
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

def adjusted1_trends(output_folder='adjusted1_plots'):
    os.makedirs(output_folder, exist_ok=True)

    # Read the Excel file
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")

    # Define prefixes and colors
    prefixes = ['adjusted1']
    colors = ['green']

    # Get list of years from column names
    pattern = re.compile(r'^adjusted1_(\d{4})$')
    years = sorted(int(pattern.match(col).group(1)) for col in df.columns if pattern.match(col))

    # Loop through each district (adm1 = FIRST_DNAM)
    for district in df['FIRST_DNAM'].dropna().unique():
        df_district = df[df['FIRST_DNAM'] == district]
        chiefdoms = df_district['FIRST_CHIE'].dropna().unique()
        n = len(chiefdoms)

        n_cols = 4
        n_rows = int(np.ceil(n / n_cols))

        fig, axes = plt.subplots(n_rows, n_cols, figsize=(n_cols * 5, n_rows * 4), sharex=False, sharey=True)
        axes = axes.flatten()

        for i, chiefdom in enumerate(chiefdoms):
            ax = axes[i]
            row = df_district[df_district['FIRST_CHIE'] == chiefdom]

            if row.empty:
                ax.set_title(f"{chiefdom} (No data)")
                ax.axis("off")
                continue

            for prefix, color in zip(prefixes, colors):
                cols = [f"{prefix}_{year}" for year in years if f"{prefix}_{year}" in row.columns]
                values = row[cols].values.flatten()

                if len(values) != len(years) or all(pd.isna(values)):
                    continue  # Skip if no valid data

                ax.plot(years, values, marker='o', label=prefix.replace('_', ' ').title(), color=color)

                # Add trend line if enough data
                if np.count_nonzero(~np.isnan(values)) >= 2:
                    fit = np.polyfit(years, values, 1)
                    trend_line = np.poly1d(fit)(years)
                    ax.plot(years, trend_line, linestyle='--', color=color, alpha=0.7, label="Trend")

            ax.set_title(chiefdom, fontsize=10)
            ax.set_xlabel("Year")  # Ensure each plot shows year
            ax.grid(True)
            ax.tick_params(axis='x', rotation=45)

        # Turn off unused subplots
        for j in range(i + 1, len(axes)):
            axes[j].axis("off")

        # Add shared legend
        handles, labels = axes[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, title="Indicator", loc="lower center", ncol=4)

        fig.suptitle(f"Adjusted1 Incidence Trends by Chiefdom - {district}", fontsize=14)
        plt.tight_layout(rect=[0, 0.05, 1, 0.95])

        filename = os.path.join(output_folder, f"{district}_trends.png")
        plt.savefig(filename, dpi=300)
        plt.close()
        print(f"[Saved] {filename}")

    return df

# Adjusted2 

def adjusted2_trends(output_folder='adjusted2_plots'):
    os.makedirs(output_folder, exist_ok=True)

    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    prefixes = ['adjusted2']
    colors = ['orange']

    pattern = re.compile(r'^adjusted2_(\d{4})$')
    years = sorted(int(pattern.match(col).group(1)) for col in df.columns if pattern.match(col))

    for district in df['FIRST_DNAM'].dropna().unique():
        df_district = df[df['FIRST_DNAM'] == district]
        chiefdoms = df_district['FIRST_CHIE'].dropna().unique()
        n = len(chiefdoms)

        n_cols = 4
        n_rows = int(np.ceil(n / n_cols))

        fig, axes = plt.subplots(n_rows, n_cols, figsize=(n_cols * 5, n_rows * 4), sharex=False, sharey=True)
        axes = axes.flatten()

        for i, chiefdom in enumerate(chiefdoms):
            ax = axes[i]
            row = df_district[df_district['FIRST_CHIE'] == chiefdom]

            if row.empty:
                ax.set_title(f"{chiefdom} (No data)")
                ax.axis("off")
                continue

            for prefix, color in zip(prefixes, colors):
                cols = [f"{prefix}_{year}" for year in years if f"{prefix}_{year}" in row.columns]
                values = row[cols].values.flatten()

                if len(values) != len(years) or all(pd.isna(values)):
                    continue

                ax.plot(years, values, marker='o', label=prefix.replace('_', ' ').title(), color=color)

                if np.count_nonzero(~np.isnan(values)) >= 2:
                    fit = np.polyfit(years, values, 1)
                    trend_line = np.poly1d(fit)(years)
                    ax.plot(years, trend_line, linestyle='--', color=color, alpha=0.7, label="Trend")

            ax.set_title(chiefdom, fontsize=10)
            ax.set_xlabel("Year")
            ax.grid(True)
            ax.tick_params(axis='x', rotation=45)

        for j in range(i + 1, len(axes)):
            axes[j].axis("off")

        handles, labels = axes[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, title="Indicator", loc="lower center", ncol=4)

        fig.suptitle(f"Adjusted2 Incidence Trends by Chiefdom - {district}", fontsize=14)
        plt.tight_layout(rect=[0, 0.05, 1, 0.95])

        filename = os.path.join(output_folder, f"{district}_trends.png")
        plt.savefig(filename, dpi=300)
        plt.close()
        print(f"[Saved] {filename}")

    return df


### Adjusted3
def adjusted3_trends(output_folder='adjusted3_plots'):
    os.makedirs(output_folder, exist_ok=True)

    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    prefixes = ['adjusted3']
    colors = ['purple']

    pattern = re.compile(r'^adjusted3_(\d{4})$')
    years = sorted(int(pattern.match(col).group(1)) for col in df.columns if pattern.match(col))

    for district in df['FIRST_DNAM'].dropna().unique():
        df_district = df[df['FIRST_DNAM'] == district]
        chiefdoms = df_district['FIRST_CHIE'].dropna().unique()
        n = len(chiefdoms)

        n_cols = 4
        n_rows = int(np.ceil(n / n_cols))

        fig, axes = plt.subplots(n_rows, n_cols, figsize=(n_cols * 5, n_rows * 4), sharex=False, sharey=True)
        axes = axes.flatten()

        for i, chiefdom in enumerate(chiefdoms):
            ax = axes[i]
            row = df_district[df_district['FIRST_CHIE'] == chiefdom]

            if row.empty:
                ax.set_title(f"{chiefdom} (No data)")
                ax.axis("off")
                continue

            for prefix, color in zip(prefixes, colors):
                cols = [f"{prefix}_{year}" for year in years if f"{prefix}_{year}" in row.columns]
                values = row[cols].values.flatten()

                if len(values) != len(years) or all(pd.isna(values)):
                    continue

                ax.plot(years, values, marker='o', label=prefix.replace('_', ' ').title(), color=color)

                if np.count_nonzero(~np.isnan(values)) >= 2:
                    fit = np.polyfit(years, values, 1)
                    trend_line = np.poly1d(fit)(years)
                    ax.plot(years, trend_line, linestyle='--', color=color, alpha=0.7, label="Trend")

            ax.set_title(chiefdom, fontsize=10)
            ax.set_xlabel("Year")
            ax.grid(True)
            ax.tick_params(axis='x', rotation=45)

        for j in range(i + 1, len(axes)):
            axes[j].axis("off")

        handles, labels = axes[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, title="Indicator", loc="lower center", ncol=4)

        fig.suptitle(f"Adjusted3 Incidence Trends by Chiefdom - {district}", fontsize=14)
        plt.tight_layout(rect=[0, 0.05, 1, 0.95])

        filename = os.path.join(output_folder, f"{district}_trends.png")
        plt.savefig(filename, dpi=300)
        plt.close()
        print(f"[Saved] {filename}")

    return df


### National crude trends
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import re

def plot_national_crude_trend(output_path='national_crude_incidence_trend.png'):
    # Identify crude incidence columns
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    pattern = re.compile(r'^crude_incidence_(\d{4})$')
    #year_cols = [col for col in df.columns if pattern.match(col)]
    year_cols = [col for col in df.columns 
                 if pattern.match(col) and 2021 <= int(pattern.match(col).group(1)) <= 2024]

    # Compute national averages per year
    averages = df[year_cols].mean(axis=0)
    avg_df = averages.reset_index()
    avg_df.columns = ['Year', 'National_Crude_Incidence']
    avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
    avg_df = avg_df.sort_values('Year').reset_index(drop=True)

    # Compute overall change (first to last year)
    y_start = avg_df['National_Crude_Incidence'].iloc[0]
    y_end = avg_df['National_Crude_Incidence'].iloc[-1]
    overall_change = ((y_end - y_start) / y_start) * 100
    subtitle_text = f"Change from {avg_df['Year'].iloc[0]} to {avg_df['Year'].iloc[-1]}: {overall_change:+.1f}%"

    # Plot setup
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(
        avg_df['Year'],
        avg_df['National_Crude_Incidence'],
        marker='o',
        color='darkblue',
        linewidth=2.5,
        label='API'
    )

    # Trend line
    if len(avg_df) >= 2:
        fit = np.polyfit(avg_df['Year'], avg_df['National_Crude_Incidence'], 1)
        trend_line = np.poly1d(fit)(avg_df['Year'])
        ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=2, label='Trend')

    # Annotate crude incidence values above each point (with box)
    for i, row in avg_df.iterrows():
        year = row['Year']
        value = row['National_Crude_Incidence']
        if pd.notna(value):
            ax.text(
                year, value + 2, f"{value:.1f}",
                fontsize=9, fontweight='bold',
                ha='center', va='bottom',
                bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
            )

    # Annotate % change between points (with box)
    for i in range(1, len(avg_df)):
        y1 = avg_df.loc[i - 1, 'National_Crude_Incidence']
        y2 = avg_df.loc[i, 'National_Crude_Incidence']
        x1 = avg_df.loc[i - 1, 'Year']
        x2 = avg_df.loc[i, 'Year']

        if pd.notna(y1) and pd.notna(y2):
            x_mid = (x1 + x2) / 2
            y_mid = (y1 + y2) / 2
            pct_change = ((y2 - y1) / y1) * 100
            label = f"{pct_change:+.0f}%"

            ax.text(
                x_mid, y_mid + 2, label,
                fontsize=8, fontweight='bold',
                ha='center', va='bottom',
                bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
            )

    # X-axis
    ax.set_xticks(list(range(2021, 2025)))
    ax.set_xlim(2021, 2024)
    ax.tick_params(axis='x', labelsize=9)
    ax.tick_params(axis='y', labelsize=9)

    # Y-axis: dynamic min, max, and step
    y_values = avg_df['National_Crude_Incidence']
    y_min = np.floor(y_values.min() / 5) * 5
    y_max = np.ceil(y_values.max() / 5) * 5
    y_range = y_max - y_min

    if y_range <= 25:
        step = 5
    elif y_range <= 50:
        step = 10
    else:
        step = 20

    ax.set_yticks(np.arange(y_min, y_max + step, step))
    ax.set_ylim(y_min, y_max + step)

    # Main title
    main_title = "Annual Parasite Incidence Trend (2021–2024)"
    subtitle_text = f"Change from {avg_df['Year'].iloc[0]} to {avg_df['Year'].iloc[-1]}: {overall_change:+.1f}%"
    
    # Set main title
    ax.set_title(
        main_title,
        fontsize=12, fontweight='bold', pad=10, loc='center'
    )
    ax.set_xlabel("Year", fontsize=10, fontweight='bold')
    ax.set_ylabel("Annual Parasite Incidence", fontsize=10, fontweight='bold')
    ax.legend(fontsize=9)
    
    # Set subtitle below the main title, inside the plot
    ax.text(
        0.5, 0.95, subtitle_text,  # slightly above the plot, below the main title
        transform=ax.transAxes,
        fontsize=9, fontweight='bold',
        ha='center', va='bottom',
        bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
    )

    # Final layout
    plt.tight_layout(rect=[0, 0, 1, 0.90])
    plt.savefig(output_path, dpi=400, bbox_inches='tight')
    plt.close()
    print(f"[Saved] {output_path}")

def plot_national_crude_trend_by_first_dnam(output_dir='plots/'):
    import os
    import matplotlib.pyplot as plt
    import pandas as pd
    import numpy as np
    import re
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Read data
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    
    # Identify crude incidence columns for 2021-2024
    pattern = re.compile(r'^crude_incidence_(\d{4})$')
    year_cols = [col for col in df.columns 
                 if pattern.match(col) and 2021 <= int(pattern.match(col).group(1)) <= 2024]
    
    # Get unique FIRST_DNAM values
    first_dnam_values = df['FIRST_DNAM'].dropna().unique()
    
    # Define colors for each FIRST_DNAM using tab10 for better distinction
    colors = plt.cm.tab10(np.linspace(0, 1, min(len(first_dnam_values), 10)))
    # If more than 10 categories, cycle through tab10 and then use tab20
    if len(first_dnam_values) > 10:
        colors2 = plt.cm.tab20(np.linspace(0, 1, len(first_dnam_values) - 10))
        colors = np.concatenate([colors, colors2])
    color_map = dict(zip(first_dnam_values, colors))
    
    # Store data for combined plot
    combined_data = {}
    
    # Create individual plots for each FIRST_DNAM
    for first_dnam in first_dnam_values:
        # Filter data for this FIRST_DNAM
        dnam_df = df[df['FIRST_DNAM'] == first_dnam]
        
        if len(dnam_df) == 0:
            continue
            
        # Compute averages per year for this FIRST_DNAM
        averages = dnam_df[year_cols].mean(axis=0)
        avg_df = averages.reset_index()
        avg_df.columns = ['Year', 'Crude_Incidence']
        avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
        avg_df = avg_df.sort_values('Year').reset_index(drop=True)
        
        # Store for combined plot
        combined_data[first_dnam] = avg_df.copy()
        
        # Skip if no valid data
        if avg_df['Crude_Incidence'].isna().all():
            continue
            
        # Compute overall change (first to last year)
        valid_data = avg_df.dropna(subset=['Crude_Incidence'])
        if len(valid_data) >= 2:
            y_start = valid_data['Crude_Incidence'].iloc[0]
            y_end = valid_data['Crude_Incidence'].iloc[-1]
            overall_change = ((y_end - y_start) / y_start) * 100
            subtitle_text = f"Change from {valid_data['Year'].iloc[0]} to {valid_data['Year'].iloc[-1]}: {overall_change:+.1f}%"
        else:
            subtitle_text = "Insufficient data for trend calculation"
        
        # Create individual plot
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Plot line with markers
        ax.plot(
            avg_df['Year'],
            avg_df['Crude_Incidence'],
            marker='o',
            color=color_map[first_dnam],
            linewidth=2.5,
            markersize=8,
            label=first_dnam
        )
        
        # Add trend line if we have enough data points
        valid_points = avg_df.dropna(subset=['Crude_Incidence'])
        if len(valid_points) >= 2:
            fit = np.polyfit(valid_points['Year'], valid_points['Crude_Incidence'], 1)
            trend_line = np.poly1d(fit)(avg_df['Year'])
            ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=2, label='Trend')
        
        # Annotate crude incidence values above each point
        for i, row in avg_df.iterrows():
            year = row['Year']
            value = row['Crude_Incidence']
            if pd.notna(value):
                ax.text(
                    year, value + 2, f"{value:.1f}",
                    fontsize=9, fontweight='bold',
                    ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
                )
        
        # Annotate % change between consecutive points
        for i in range(1, len(avg_df)):
            y1 = avg_df.loc[i - 1, 'Crude_Incidence']
            y2 = avg_df.loc[i, 'Crude_Incidence']
            x1 = avg_df.loc[i - 1, 'Year']
            x2 = avg_df.loc[i, 'Year']
            
            if pd.notna(y1) and pd.notna(y2):
                x_mid = (x1 + x2) / 2
                y_mid = (y1 + y2) / 2
                pct_change = ((y2 - y1) / y1) * 100
                label = f"{pct_change:+.0f}%"
                
                ax.text(
                    x_mid, y_mid + 2, label,
                    fontsize=8, fontweight='bold',
                    ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
                )
        
        # Format axes
        ax.set_xticks(list(range(2021, 2025)))
        ax.set_xlim(2020.5, 2024.5)
        ax.tick_params(axis='x', labelsize=9)
        ax.tick_params(axis='y', labelsize=9)
        
        # Dynamic Y-axis
        y_values = avg_df['Crude_Incidence'].dropna()
        if len(y_values) > 0:
            y_min = np.floor(y_values.min() / 5) * 5
            y_max = np.ceil(y_values.max() / 5) * 5
            y_range = y_max - y_min
            
            if y_range <= 25:
                step = 5
            elif y_range <= 50:
                step = 10
            else:
                step = 20
                
            ax.set_yticks(np.arange(y_min, y_max + step, step))
            ax.set_ylim(y_min, y_max + step)  # Reduced padding since text is closer
        
        # Titles and labels
        main_title = f"Annual Parasite Incidence Trend - {first_dnam} (2021–2024)"
        ax.set_title(main_title, fontsize=12, fontweight='bold', pad=15)
        ax.set_xlabel("Year", fontsize=10, fontweight='bold')
        ax.set_ylabel("Annual Parasite Incidence", fontsize=10, fontweight='bold')
        ax.legend(fontsize=9)
        
        # Add subtitle
        ax.text(
            0.5, 0.95, subtitle_text,
            transform=ax.transAxes,
            fontsize=9, fontweight='bold',
            ha='center', va='bottom',
            bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3')
        )
        
        # Save individual plot
        safe_filename = "".join(c for c in first_dnam if c.isalnum() or c in (' ', '-', '_')).rstrip()
        output_path = os.path.join(output_dir, f'crude_incidence_trend_{safe_filename}.png')
        plt.tight_layout()
        plt.savefig(output_path, dpi=400, bbox_inches='tight')
        plt.close()
        print(f"[Saved] {output_path}")
    

    
    # Create subplot version with 4x4 grid
    n_plots = len([data for data in combined_data.values() if not data['Crude_Incidence'].isna().all()])
    n_rows = 4
    n_cols = 4
    
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(20, 16))
    fig.suptitle("Annual Parasite Incidence Trend by FIRST_DNAM (2021–2024)", 
                 fontsize=16, fontweight='bold', y=0.98)
    
    # Flatten axes for easier iteration
    axes_flat = axes.flatten()
    
    plot_idx = 0
    for first_dnam, avg_df in combined_data.items():
        if avg_df['Crude_Incidence'].isna().all() or plot_idx >= n_rows * n_cols:
            continue
            
        ax = axes_flat[plot_idx]
        
        # Calculate overall change for title
        valid_data = avg_df.dropna(subset=['Crude_Incidence'])
        if len(valid_data) >= 2:
            y_start = valid_data['Crude_Incidence'].iloc[0]
            y_end = valid_data['Crude_Incidence'].iloc[-1]
            overall_change = ((y_end - y_start) / y_start) * 100
            title_text = f"{first_dnam}\nOverall: {overall_change:+.1f}%"
        else:
            title_text = f"{first_dnam}\nInsufficient data"
        
        # Plot the data with unique color
        ax.plot(
            avg_df['Year'],
            avg_df['Crude_Incidence'],
            marker='o',
            color=color_map[first_dnam],
            linewidth=2,
            markersize=4,
            label=first_dnam
        )
        
        # Add trend line if we have enough data points
        valid_points = avg_df.dropna(subset=['Crude_Incidence'])
        if len(valid_points) >= 2:
            fit = np.polyfit(valid_points['Year'], valid_points['Crude_Incidence'], 1)
            trend_line = np.poly1d(fit)(avg_df['Year'])
            ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=1, alpha=0.7)
        
        # Add percentage change annotations between consecutive points
        for i in range(1, len(avg_df)):
            y1 = avg_df.loc[i - 1, 'Crude_Incidence']
            y2 = avg_df.loc[i, 'Crude_Incidence']
            x1 = avg_df.loc[i - 1, 'Year']
            x2 = avg_df.loc[i, 'Year']
            
            if pd.notna(y1) and pd.notna(y2):
                x_mid = (x1 + x2) / 2
                y_mid = (y1 + y2) / 2
                pct_change = ((y2 - y1) / y1) * 100
                label = f"{pct_change:+.0f}%"
                
                ax.text(
                    x_mid, y_mid + 1, label,
                    fontsize=7, fontweight='bold',
                    ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.2')
                )
        
        # Format subplot
        ax.set_title(title_text, fontsize=9, fontweight='bold', pad=5)
        ax.set_xticks([2021, 2022, 2023, 2024])
        ax.tick_params(axis='x', labelsize=8)
        ax.tick_params(axis='y', labelsize=8)
        ax.grid(True, alpha=0.3)
        
        # Set y-axis limits based on data
        y_values = avg_df['Crude_Incidence'].dropna()
        if len(y_values) > 0:
            y_min = max(0, y_values.min() - 2)
            y_max = y_values.max() + 3  # Extra space for annotations
            ax.set_ylim(y_min, y_max)
        
        plot_idx += 1
    
    # Hide unused subplots
    for i in range(plot_idx, len(axes_flat)):
        axes_flat[i].set_visible(False)
    
    # Add common labels
    fig.text(0.5, 0.02, 'Year', ha='center', va='center', fontsize=12, fontweight='bold')
    fig.text(0.02, 0.5, 'Annual Parasite Incidence', ha='center', va='center', 
             rotation='vertical', fontsize=12, fontweight='bold')
    
    # Save subplot version
    subplot_output_path = os.path.join(output_dir, 'crude_incidence_trend_subplots_4x4.png')
    plt.tight_layout(rect=[0.03, 0.03, 1, 0.96])
    plt.savefig(subplot_output_path, dpi=400, bbox_inches='tight')
    plt.close()
    print(f"[Saved] {subplot_output_path}")
    
    print(f"\nCreated {len(first_dnam_values)} individual plots and 1 subplot grid")
    print(f"All plots saved in: {output_dir}")
    
    # Create subplot plots for each FIRST_DNAM showing all FIRST_CHIE (4 columns, n rows)
    for first_dnam in first_dnam_values:
        # Filter data for this FIRST_DNAM
        dnam_df = df[df['FIRST_DNAM'] == first_dnam]
        
        if len(dnam_df) == 0:
            continue
            
        # Get unique FIRST_CHIE values for this FIRST_DNAM
        first_chie_values = dnam_df['FIRST_CHIE'].dropna().unique()
        
        if len(first_chie_values) == 0:
            continue
        
        # Define colors for FIRST_CHIE within this FIRST_DNAM
        chie_colors = plt.cm.tab10(np.linspace(0, 1, min(len(first_chie_values), 10)))
        if len(first_chie_values) > 10:
            chie_colors2 = plt.cm.tab20(np.linspace(0, 1, len(first_chie_values) - 10))
            chie_colors = np.concatenate([chie_colors, chie_colors2])
        chie_color_map = dict(zip(first_chie_values, chie_colors))
        
        # Calculate grid dimensions (4 columns, n rows)
        n_cols = 4
        n_rows = int(np.ceil(len(first_chie_values) / n_cols))
        
        # Create figure for this FIRST_DNAM
        fig, axes = plt.subplots(n_rows, n_cols, figsize=(20, 5 * n_rows))
        fig.suptitle(f"Annual Parasite Incidence Trend by FIRST_CHIE - {first_dnam} (2021–2024)", 
                     fontsize=16, fontweight='bold', y=0.98)
        
        # Handle single row case
        if n_rows == 1:
            axes = axes.reshape(1, -1) if len(first_chie_values) > 1 else [axes]
        
        # Flatten axes for easier iteration
        axes_flat = axes.flatten()
        
        plot_idx = 0
        for first_chie in first_chie_values:
            if plot_idx >= len(axes_flat):
                break
                
            # Filter data for this FIRST_CHIE
            chie_df = dnam_df[dnam_df['FIRST_CHIE'] == first_chie]
            
            if len(chie_df) == 0:
                plot_idx += 1
                continue
                
            # Compute averages per year for this FIRST_CHIE
            averages = chie_df[year_cols].mean(axis=0)
            avg_df = averages.reset_index()
            avg_df.columns = ['Year', 'Crude_Incidence']
            avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
            avg_df = avg_df.sort_values('Year').reset_index(drop=True)
            
            ax = axes_flat[plot_idx]
            
            # Calculate overall change for title
            valid_data = avg_df.dropna(subset=['Crude_Incidence'])
            if len(valid_data) >= 2:
                y_start = valid_data['Crude_Incidence'].iloc[0]
                y_end = valid_data['Crude_Incidence'].iloc[-1]
                overall_change = ((y_end - y_start) / y_start) * 100
                title_text = f"{first_chie}\nOverall: {overall_change:+.1f}%"
            else:
                title_text = f"{first_chie}\nInsufficient data"
            
            # Plot the data with unique color for this FIRST_CHIE
            ax.plot(
                avg_df['Year'],
                avg_df['Crude_Incidence'],
                marker='o',
                color=chie_color_map[first_chie],
                linewidth=2,
                markersize=4,
                label=first_chie
            )
            
            # Add trend line if we have enough data points
            valid_points = avg_df.dropna(subset=['Crude_Incidence'])
            if len(valid_points) >= 2:
                fit = np.polyfit(valid_points['Year'], valid_points['Crude_Incidence'], 1)
                trend_line = np.poly1d(fit)(avg_df['Year'])
                ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=1, alpha=0.7)
            
            # Add percentage change annotations between consecutive points
            for i in range(1, len(avg_df)):
                y1 = avg_df.loc[i - 1, 'Crude_Incidence']
                y2 = avg_df.loc[i, 'Crude_Incidence']
                x1 = avg_df.loc[i - 1, 'Year']
                x2 = avg_df.loc[i, 'Year']
                
                if pd.notna(y1) and pd.notna(y2):
                    x_mid = (x1 + x2) / 2
                    y_mid = (y1 + y2) / 2
                    pct_change = ((y2 - y1) / y1) * 100
                    label = f"{pct_change:+.0f}%"
                    
                    ax.text(
                        x_mid, y_mid + 1, label,
                        fontsize=7, fontweight='bold',
                        ha='center', va='bottom',
                        bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.2')
                    )
            
            # Format subplot
            ax.set_title(title_text, fontsize=9, fontweight='bold', pad=5)
            ax.set_xticks([2021, 2022, 2023, 2024])
            ax.tick_params(axis='x', labelsize=8)
            ax.tick_params(axis='y', labelsize=8)
            ax.grid(True, alpha=0.3)
            
            # Set y-axis limits based on data
            y_values = avg_df['Crude_Incidence'].dropna()
            if len(y_values) > 0:
                y_min = max(0, y_values.min() - 2)
                y_max = y_values.max() + 3  # Extra space for annotations
                ax.set_ylim(y_min, y_max)
            
            plot_idx += 1
        
        # Hide unused subplots
        for i in range(plot_idx, len(axes_flat)):
            axes_flat[i].set_visible(False)
        
        # Add common labels
        fig.text(0.5, 0.02, 'Year', ha='center', va='center', fontsize=12, fontweight='bold')
        fig.text(0.02, 0.5, 'Annual Parasite Incidence', ha='center', va='center', 
                 rotation='vertical', fontsize=12, fontweight='bold')
        
        # Save subplot version for this FIRST_DNAM
        safe_dnam_name = "".join(c for c in first_dnam if c.isalnum() or c in (' ', '-', '_')).rstrip()
        chie_subplot_path = os.path.join(output_dir, f'crude_incidence_CHIE_subplots_{safe_dnam_name}.png')
        plt.tight_layout(rect=[0.03, 0.03, 1, 0.96])
        plt.savefig(chie_subplot_path, dpi=400, bbox_inches='tight')
        plt.close()
        print(f"[Saved] {chie_subplot_path}")
    
    print(f"\nCreated {len(first_dnam_values)} individual plots, 1 DNAM subplot grid, and {len(first_dnam_values)} CHIE subplot grids")
    print(f"All plots saved in: {output_dir}")


### National Adjusted1
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import re


def plot_national_adjusted1_trend(output_path='national_adjusted1_incidence_trend.png'):
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    pattern = re.compile(r'^adjusted1_(\d{4})$')
    year_cols = [col for col in df.columns if pattern.match(col)]

    averages = df[year_cols].mean(axis=0)
    avg_df = averages.reset_index()
    avg_df.columns = ['Year', 'National_Adjusted1_Incidence']
    avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
    avg_df = avg_df.sort_values('Year').reset_index(drop=True)

    y_start = avg_df['National_Adjusted1_Incidence'].iloc[0]
    y_end = avg_df['National_Adjusted1_Incidence'].iloc[-1]
    overall_change = ((y_end - y_start) / y_start) * 100
    subtitle_text = f"Change from {avg_df['Year'].iloc[0]} to {avg_df['Year'].iloc[-1]}: {overall_change:+.1f}%"

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(avg_df['Year'], avg_df['National_Adjusted1_Incidence'], marker='o', color='darkblue', linewidth=2.5, label='API')

    if len(avg_df) >= 2:
        fit = np.polyfit(avg_df['Year'], avg_df['National_Adjusted1_Incidence'], 1)
        trend_line = np.poly1d(fit)(avg_df['Year'])
        ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=2, label='Trend')

    for i, row in avg_df.iterrows():
        if pd.notna(row['National_Adjusted1_Incidence']):
            ax.text(row['Year'], row['National_Adjusted1_Incidence'] + 2, f"{row['National_Adjusted1_Incidence']:.1f}",
                    fontsize=9, fontweight='bold', ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    for i in range(1, len(avg_df)):
        y1, y2 = avg_df.loc[i - 1, 'National_Adjusted1_Incidence'], avg_df.loc[i, 'National_Adjusted1_Incidence']
        x1, x2 = avg_df.loc[i - 1, 'Year'], avg_df.loc[i, 'Year']
        if pd.notna(y1) and pd.notna(y2):
            x_mid = (x1 + x2) / 2
            y_mid = (y1 + y2) / 2
            pct_change = ((y2 - y1) / y1) * 100
            ax.text(x_mid, y_mid + 2, f"{pct_change:+.0f}%", fontsize=8, fontweight='bold', ha='center',
                    va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    ax.set_xticks(list(range(2015, 2025)))
    ax.set_xlim(2015, 2024)
    ax.tick_params(axis='x', labelsize=9)
    ax.tick_params(axis='y', labelsize=9)

    y_values = avg_df['National_Adjusted1_Incidence']
    y_min, y_max = np.floor(y_values.min() / 5) * 5, np.ceil(y_values.max() / 5) * 5
    step = 5 if (y_max - y_min) <= 25 else (10 if (y_max - y_min) <= 50 else 20)

    ax.set_yticks(np.arange(y_min, y_max + step, step))
    ax.set_ylim(y_min, y_max + step)

    ax.set_title("Adjusted1 Incidence Trend (2015–2024)", fontsize=12, fontweight='bold', pad=10)
    ax.set_xlabel("Year", fontsize=10, fontweight='bold')
    ax.set_ylabel("Adjusted1 Incidence", fontsize=10, fontweight='bold')
    ax.legend(fontsize=9)
    ax.text(0.5, 0.95, subtitle_text, transform=ax.transAxes, fontsize=9, fontweight='bold',
            ha='center', va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    plt.tight_layout(rect=[0, 0, 1, 0.90])
    plt.savefig(output_path, dpi=400, bbox_inches='tight')
    plt.close()
    print(f"[Saved] {output_path}")


### National adjusted2
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import re

def plot_national_adjusted2_trend(output_path='national_adjusted2_incidence_trend.png'):
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    pattern = re.compile(r'^adjusted2_(\d{4})$')
    year_cols = [col for col in df.columns if pattern.match(col)]

    averages = df[year_cols].mean(axis=0)
    avg_df = averages.reset_index()
    avg_df.columns = ['Year', 'National_Adjusted2_Incidence']
    avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
    avg_df = avg_df.sort_values('Year').reset_index(drop=True)

    y_start = avg_df['National_Adjusted2_Incidence'].iloc[0]
    y_end = avg_df['National_Adjusted2_Incidence'].iloc[-1]
    overall_change = ((y_end - y_start) / y_start) * 100
    subtitle_text = f"Change from {avg_df['Year'].iloc[0]} to {avg_df['Year'].iloc[-1]}: {overall_change:+.1f}%"

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(avg_df['Year'], avg_df['National_Adjusted2_Incidence'], marker='o', color='darkgreen', linewidth=2.5, label='API')

    if len(avg_df) >= 2:
        fit = np.polyfit(avg_df['Year'], avg_df['National_Adjusted2_Incidence'], 1)
        trend_line = np.poly1d(fit)(avg_df['Year'])
        ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=2, label='Trend')

    for i, row in avg_df.iterrows():
        if pd.notna(row['National_Adjusted2_Incidence']):
            ax.text(row['Year'], row['National_Adjusted2_Incidence'] + 2, f"{row['National_Adjusted2_Incidence']:.1f}",
                    fontsize=9, fontweight='bold', ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    for i in range(1, len(avg_df)):
        y1, y2 = avg_df.loc[i - 1, 'National_Adjusted2_Incidence'], avg_df.loc[i, 'National_Adjusted2_Incidence']
        x1, x2 = avg_df.loc[i - 1, 'Year'], avg_df.loc[i, 'Year']
        if pd.notna(y1) and pd.notna(y2):
            x_mid = (x1 + x2) / 2
            y_mid = (y1 + y2) / 2
            pct_change = ((y2 - y1) / y1) * 100
            ax.text(x_mid, y_mid + 2, f"{pct_change:+.0f}%", fontsize=8, fontweight='bold', ha='center',
                    va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    ax.set_xticks(list(range(2015, 2025)))
    ax.set_xlim(2015, 2024)
    ax.tick_params(axis='x', labelsize=9)
    ax.tick_params(axis='y', labelsize=9)

    y_values = avg_df['National_Adjusted2_Incidence']
    y_min, y_max = np.floor(y_values.min() / 5) * 5, np.ceil(y_values.max() / 5) * 5
    step = 5 if (y_max - y_min) <= 25 else (10 if (y_max - y_min) <= 50 else 20)

    ax.set_yticks(np.arange(y_min, y_max + step, step))
    ax.set_ylim(y_min, y_max + step)

    ax.set_title("Adjusted2 Incidence Trend (2015–2024)", fontsize=12, fontweight='bold', pad=10)
    ax.set_xlabel("Year", fontsize=10, fontweight='bold')
    ax.set_ylabel("Adjusted2 Incidence", fontsize=10, fontweight='bold')
    ax.legend(fontsize=9)
    ax.text(0.5, 0.95, subtitle_text, transform=ax.transAxes, fontsize=9, fontweight='bold',
            ha='center', va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    plt.tight_layout(rect=[0, 0, 1, 0.90])
    plt.savefig(output_path, dpi=400, bbox_inches='tight')
    plt.close()
    print(f"[Saved] {output_path}")


### National adjusted3

def plot_national_adjusted3_trend(output_path='national_adjusted3_incidence_trend.png'):
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    pattern = re.compile(r'^adjusted3_(\d{4})$')
    year_cols = [col for col in df.columns if pattern.match(col)]

    averages = df[year_cols].mean(axis=0)
    avg_df = averages.reset_index()
    avg_df.columns = ['Year', 'National_Adjusted3_Incidence']
    avg_df['Year'] = avg_df['Year'].str.extract(r'(\d{4})').astype(int)
    avg_df = avg_df.sort_values('Year').reset_index(drop=True)

    y_start = avg_df['National_Adjusted3_Incidence'].iloc[0]
    y_end = avg_df['National_Adjusted3_Incidence'].iloc[-1]
    overall_change = ((y_end - y_start) / y_start) * 100
    subtitle_text = f"Change from {avg_df['Year'].iloc[0]} to {avg_df['Year'].iloc[-1]}: {overall_change:+.1f}%"

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(avg_df['Year'], avg_df['National_Adjusted3_Incidence'], marker='o', color='darkred', linewidth=2.5, label='API')

    if len(avg_df) >= 2:
        fit = np.polyfit(avg_df['Year'], avg_df['National_Adjusted3_Incidence'], 1)
        trend_line = np.poly1d(fit)(avg_df['Year'])
        ax.plot(avg_df['Year'], trend_line, linestyle='--', color='gray', linewidth=2, label='Trend')

    for i, row in avg_df.iterrows():
        if pd.notna(row['National_Adjusted3_Incidence']):
            ax.text(row['Year'], row['National_Adjusted3_Incidence'] + 2, f"{row['National_Adjusted3_Incidence']:.1f}",
                    fontsize=9, fontweight='bold', ha='center', va='bottom',
                    bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    for i in range(1, len(avg_df)):
        y1, y2 = avg_df.loc[i - 1, 'National_Adjusted3_Incidence'], avg_df.loc[i, 'National_Adjusted3_Incidence']
        x1, x2 = avg_df.loc[i - 1, 'Year'], avg_df.loc[i, 'Year']
        if pd.notna(y1) and pd.notna(y2):
            x_mid = (x1 + x2) / 2
            y_mid = (y1 + y2) / 2
            pct_change = ((y2 - y1) / y1) * 100
            ax.text(x_mid, y_mid + 2, f"{pct_change:+.0f}%", fontsize=8, fontweight='bold', ha='center',
                    va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    ax.set_xticks(list(range(2015, 2025)))
    ax.set_xlim(2015, 2024)
    ax.tick_params(axis='x', labelsize=9)
    ax.tick_params(axis='y', labelsize=9)

    y_values = avg_df['National_Adjusted3_Incidence']
    y_min, y_max = np.floor(y_values.min() / 5) * 5, np.ceil(y_values.max() / 5) * 5
    step = 5 if (y_max - y_min) <= 25 else (10 if (y_max - y_min) <= 50 else 20)

    ax.set_yticks(np.arange(y_min, y_max + step, step))
    ax.set_ylim(y_min, y_max + step)

    ax.set_title("Adjusted3 Incidence Trend (2015–2024)", fontsize=12, fontweight='bold', pad=10)
    ax.set_xlabel("Year", fontsize=10, fontweight='bold')
    ax.set_ylabel("Adjusted3 Incidence", fontsize=10, fontweight='bold')
    ax.legend(fontsize=9)
    ax.text(0.5, 0.95, subtitle_text, transform=ax.transAxes, fontsize=9, fontweight='bold',
            ha='center', va='bottom', bbox=dict(facecolor='white', edgecolor='black', boxstyle='round,pad=0.3'))

    plt.tight_layout(rect=[0, 0, 1, 0.90])
    plt.savefig(output_path, dpi=400, bbox_inches='tight')
    plt.close()
    print(f"[Saved] {output_path}")




## Word documents
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime



def add_figure(doc, image_path, caption, fig_num):
    doc.add_page_break()
    doc.add_heading(f"Figure {fig_num}", level=2)
    doc.add_picture(image_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Figure {fig_num}: {caption}", style='Caption')

def compute_slope(values, years):
    if len(values) < 2:
        return 0
    x = np.array(years)
    y = np.array(values)
    return np.polyfit(x, y, 1)[0]

def summarize_all_district_trends(df):
    years = [int(col.split('_')[-1]) for col in df.columns if col.startswith('crude_incidence_')]
    years = sorted(set(years))
    results = []

    for district in df['FIRST_DNAM'].unique():
        chiefdoms = df[df['FIRST_DNAM'] == district]['FIRST_CHIE'].unique()

        for chiefdom in chiefdoms:
            row = df[(df['FIRST_DNAM'] == district) & (df['FIRST_CHIE'] == chiefdom)].iloc[0]

            summary = {'District': district, 'Chiefdom': chiefdom}
            for prefix in ['crude_incidence', 'adjusted1', 'adjusted2', 'adjusted3']:
                cols = [f"{prefix}_{y}" for y in years if f"{prefix}_{y}" in row]
                values = [row[c] for c in cols]
                slope = compute_slope(values, years)
                trend = (
                    "increasing" if slope > 5 else
                    "decreasing" if slope < -5 else
                    "stable"
                )
                summary[prefix] = trend
            results.append(summary)

    return pd.DataFrame(results)

def interpret_district_trends(summary_df, district_name):
    output = [f"District: {district_name}"]
    df = summary_df[summary_df['District'] == district_name]
    for prefix in ['crude_incidence', 'adjusted1', 'adjusted2', 'adjusted3']:
        counts = df[prefix].value_counts()
        increasing = counts.get('increasing', 0)
        decreasing = counts.get('decreasing', 0)
        stable = counts.get('stable', 0)

        statement = (
            f"{prefix.replace('_', ' ').title()}: "
            f"{increasing} increasing, {decreasing} decreasing, {stable} stable."
        )

        if increasing > decreasing:
            statement += (
                " Rising transmission is observed in a majority of chiefdoms, which may indicate emerging outbreaks, "
                "increased reporting, or gaps in control coverage. Urgent attention is recommended, including enhanced "
                "surveillance, mass distribution of LLINs, indoor residual spraying (IRS) in hotspot areas, and "
                "reinforcing case management capacity at facility level."
            )
        elif decreasing > increasing:
            statement += (
                " The majority of chiefdoms show declining trends, suggesting that current interventions are having a positive effect. "
                "Continued investment in malaria control should be maintained, including regular monitoring, community engagement, "
                "and ensuring commodity availability. Consider conducting impact evaluations to identify successful strategies."
            )
        else:
            statement += (
                " Trends are mixed across chiefdoms. Some areas show improvement while others worsen or remain unchanged. "
                "This heterogeneity may reflect differences in intervention coverage, health system performance, or ecological factors. "
                "A tailored approach is advised—strengthening high-burden areas and maintaining gains in others."
            )

        output.append(statement)
    return output

def add_trend_summary_table(doc, trend_df, district_name):
    doc.add_heading(f"Chiefdom-Level Trend Summary: {district_name}", level=2)
    doc.add_paragraph("This table summarizes the trend direction for each incidence indicator across chiefdoms in the district.")

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Chiefdom"
    hdr_cells[1].text = "Crude Incidence"
    hdr_cells[2].text = "Adjusted1"
    hdr_cells[3].text = "Adjusted2"
    hdr_cells[4].text = "Adjusted3"

    for _, row in trend_df[trend_df['District'] == district_name].iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Chiefdom'])
        cells[1].text = row['crude_incidence']
        cells[2].text = row['adjusted1']
        cells[3].text = row['adjusted2']
        cells[4].text = row['adjusted3']

###

def export_and_interpret(
    report_folder="final_report",
    report_title="Malaria Epidemiological Analysis Report",
    author="Malaria Surveillance Team",
    subplots_folder="subplots",
    trends_folder="epi_lineplots",
    crude_trends_folder="crude_plots",
    adjusted1_trends_folder="adjusted1_plots",
    adjusted2_trends_folder="adjusted2_plots",
    adjusted3_trends_folder="adjusted3_plots"
):
    os.makedirs(report_folder, exist_ok=True)
    epi_data = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    doc = Document()

    doc.add_heading(report_title, level=0)
    p = doc.add_paragraph()
    p.add_run(f"Prepared by: {author}").bold = True
    p.add_run(f"\nDate: {datetime.now().strftime('%B %d, %Y')}")


    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This report presents the results of malaria epidemiological analysis using routine "
        "surveillance data. The analysis includes data cleaning, outlier detection, "
        "incidence calculation with various adjustment methods, and geographic distribution "
        "visualization. This document provides interpretations and recommendations based on "
        "the findings."
    )

    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "The analysis workflow involved several steps:\n"
        "1. Data concatenation and cleaning from routine surveillance files\n"
        "2. Outlier detection using IQR method and winsorization for correction\n"
        "3. Calculation of crude and adjusted incidence rates\n"
        "4. Visualization of geographic distribution of malaria burden\n"
        "5. Statistical summary and interpretation of findings"
    )

    trend_df = summarize_all_district_trends(epi_data)

    doc.add_heading("Trend Summary by District", level=1)
    for district in trend_df['District'].unique():
        trend_summary = interpret_district_trends(trend_df, district)
        for s in trend_summary:
            doc.add_paragraph(s)
        add_trend_summary_table(doc, trend_df, district)
   # Map subplots
    fig_num = 1
    
    doc.add_heading("Spatial Distribution Maps", level=1)
    for prefix in ["crude_incidence", "adjusted1", "adjusted2", "adjusted3"]:
        subplot_path = os.path.join(subplots_folder, f"{prefix}_maps.png")
        if os.path.exists(subplot_path):
            caption = f"{prefix.replace('_', ' ').title()} spatial distribution across chiefdoms"
            add_figure(doc, subplot_path, caption, fig_num)
            fig_num += 1




    # crude trends
    doc.add_heading("Crude Incidence Trends", level=1)
    if os.path.exists(crude_trends_folder):
        for file in sorted(Path(crude_trends_folder).glob("*.png")):
            district_name = file.stem
            caption = f"Crude incidence trends in {district_name}"
            add_figure(doc, str(file), caption, fig_num)
            fig_num += 1

   # Adjusted1 trends
    doc.add_heading("Adjusted1 Incidence Trends", level=1)
    if os.path.exists(adjusted1_trends_folder):
        for file in sorted(Path(adjusted1_trends_folder).glob("*.png")):
            district_name = file.stem
            caption = f"Adjusted1 incidence trends in {district_name}"
            add_figure(doc, str(file), caption, fig_num)
            fig_num += 1


      # Adjusted2 trends
    doc.add_heading("Adjusted2 Incidence Trends", level=1)
    if os.path.exists(adjusted2_trends_folder):
        for file in sorted(Path(adjusted2_trends_folder).glob("*.png")):
            district_name = file.stem
            caption = f"Adjusted2 incidence trends in {district_name}"
            add_figure(doc, str(file), caption, fig_num)
            fig_num += 1


    # Adjusted3 trends
    doc.add_heading("Adjusted3 Incidence Trends", level=1)
    if os.path.exists(adjusted3_trends_folder):
        for file in sorted(Path(adjusted3_trends_folder).glob("*.png")):
            district_name = file.stem
            caption = f"Adjusted2 incidence trends {district_name}"
            add_figure(doc, str(file), caption, fig_num)
            fig_num += 1

    
    # Incidence trends
    doc.add_heading("Incidence Trends", level=1)
    if os.path.exists(trends_folder):
        for file in sorted(Path(trends_folder).glob("*.png")):
            district_name = file.stem
            caption = f"Incidence trends in {district_name}"
            add_figure(doc, str(file), caption, fig_num)
            fig_num += 1

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(report_folder, f"Malaria_Analysis_Report_{timestamp}.docx")
    doc.save(output_file)
    print(f"\n✅ Report saved to: {output_file}")  



#### District Chiefdom Maps

import os
import re
import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from matplotlib.colors import BoundaryNorm
from matplotlib.patches import Patch

def plot_district_chiefdom_maps(
    df, shapefile, prefix, colormap='RdYlBu_r', bins=None, bin_labels=None, output_root='epi_maps'
):
    import matplotlib.pyplot as plt
    from matplotlib.colors import BoundaryNorm
    from matplotlib.patches import Patch
    import numpy as np
    import os

    # Default bins if not specified
    if bins is None:
        bins = [0, 50, 100, 250, 450, 700, 1000, float('inf')]
    if bin_labels is None:
        bin_labels = ['<50', '50-100', '100-250', '250-450', '450-700', '700-1000', '>1000']

    # Merge
    gdf = shapefile.merge(df, on=['FIRST_DNAM', 'FIRST_CHIE'], how='left', validate='1:1')

    # Detect valid columns
    import re
    pattern = re.compile(rf"^{re.escape(prefix)}\d{{4}}$")
    columns = [col for col in gdf.columns if pattern.match(col)]
    if not columns:
        print(f"No columns found for prefix '{prefix}'")
        return gdf

    cmap = plt.cm.get_cmap(colormap, len(bins) - 1)
    norm = BoundaryNorm(bins, ncolors=cmap.N)

    for column in columns:
        year = column[-4:]
        title_prefix = column.replace("_", " ").title()
        output_folder = os.path.join(output_root, prefix.rstrip('_'))
        os.makedirs(output_folder, exist_ok=True)

        for district in gdf['FIRST_DNAM'].dropna().unique():
            gdf_district = gdf[gdf['FIRST_DNAM'] == district]
            if column not in gdf_district.columns or gdf_district[column].dropna().empty:
                continue

            gdf_district[column] = gdf_district[column].round().astype(int)

            fig, ax = plt.subplots(figsize=(8, 6), dpi=300)

            gdf_district.plot(
                column=column,
                cmap=cmap,
                norm=norm,
                edgecolor='gray',
                linewidth=0.5,
                legend=False,
                ax=ax,
                missing_kwds={'color': 'lightgrey', 'edgecolor': 'white', 'linewidth': 0.5}
            )

            gdf_district.dissolve(by="FIRST_DNAM").boundary.plot(ax=ax, color="black", linewidth=1.0)

            # Add chiefdom names
            placed_centroids = []
            for _, row in gdf_district.iterrows():
                if row['geometry'] is not None and not row['geometry'].is_empty:
                    centroid = row['geometry'].centroid
                    label = str(row['FIRST_CHIE'])
                    too_close = any(
                        np.sqrt((centroid.x - x)**2 + (centroid.y - y)**2) < 0.1
                        for x, y in placed_centroids
                    )
                    placed_centroids.append((centroid.x, centroid.y))
                    ax.text(
                        centroid.x, centroid.y, label,
                        fontsize=5, ha='center', va='center',
                        color='black',
                        rotation=45 if too_close else 0,
                        rotation_mode='anchor'
                    )

            # Add legend
            legend_elements = [
                Patch(facecolor=cmap(norm(bins[i])), edgecolor='black', label=bin_labels[i])
                for i in range(len(bins) - 1)
            ]
            ax.legend(
                handles=legend_elements,
                loc='lower center',
                bbox_to_anchor=(0.5, -0.15),
                ncol=4,
                fontsize=7,
                title="Cases per 1000",
                title_fontsize=8,
                frameon=True
            )

            ax.set_title(f"{title_prefix}\n{district}", fontsize=13, fontweight='bold', pad=10)
            ax.axis("off")
            fig.subplots_adjust(bottom=0.2)

            filename = os.path.join(output_folder, f"{district.replace(' ', '_')}_{year}.png")
            plt.savefig(filename, dpi=300, bbox_inches="tight")
            plt.close()
            print(f"Saved: {filename}")

    return gdf

def crude_incidence_district_chiefdom_map():
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    return plot_district_chiefdom_maps(df, shapefile, prefix='crude_incidence_')

def adjusted1_district_chiefdom_map():
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    return plot_district_chiefdom_maps(df, shapefile, prefix='adjusted1_')

def adjusted2_district_chiefdom_map():
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    return plot_district_chiefdom_maps(df, shapefile, prefix='adjusted2_')

def adjusted3_district_chiefdom_map():
    df = pd.read_excel("input_files/others/2024_snt_data.xlsx")
    shapefile = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    return plot_district_chiefdom_maps(df, shapefile, prefix='adjusted3_')


#### Rainfall
import geopandas as gpd
import rasterio
import rasterio.mask
import numpy as np
import os
import requests
import gzip
import shutil
import tempfile
import pandas as pd
import math
from datetime import datetime
import matplotlib.pyplot as plt
from pathlib import Path


def rainfall(start_date, end_date, output_dir=None):
    """
    Calculate mean rainfall for a given shapefile over a specified date range,
    save data to Excel, and generate maps for each time period.
    
    Parameters:
    -----------
    start_date : str
        Start date in format 'YYYY-MM'
    end_date : str
        End date in format 'YYYY-MM'
    shapefile_path : str
        Path to the shapefile (.shp). The .shx and .dbf files must be in the same directory.
    output_dir : str, optional
        Directory to save outputs. If None, uses current directory.
    
    Returns:
    --------
    pandas.DataFrame
        DataFrame containing rainfall data for each region and time period
    """
    # Validate and parse dates
    try:
        start = datetime.strptime(start_date, '%Y-%m')
        end = datetime.strptime(end_date, '%Y-%m')
        if start > end:
            raise ValueError("Start date must be before end date")
    except ValueError as e:
        if "time data" in str(e):
            raise ValueError("Dates must be in format YYYY-MM")
        else:
            raise e
    
    # Set up output directory
    if output_dir is None:
        output_dir = os.getcwd()
    output_dir = Path(output_dir)
    
    # Create directories if they don't exist
    maps_dir = output_dir / "rainfall maps"
    maps_dir.mkdir(parents=True, exist_ok=True)
    
    # Check if shapefile exists and has required components
    shp_path = Path("input_files/routine/shapefile/Chiefdom2021.shp")
    if not shp_path.exists():
        raise FileNotFoundError(f"Shapefile not found: {shapefile_path}")
    
    # Check for .shx and .dbf files
    shx_path = shp_path.with_suffix('.shx')
    dbf_path = shp_path.with_suffix('.dbf')
    
    if not shx_path.exists():
        raise FileNotFoundError(f".shx file not found: {shx_path}")
    if not dbf_path.exists():
        raise FileNotFoundError(f".dbf file not found: {dbf_path}")
    
    # Load the shapefile
    try:
        gdf = gpd.read_file("input_files/routine/shapefile/Chiefdom2021.shp")
    except Exception as e:
        raise ValueError(f"Error loading shapefile: {str(e)}")
    
    # Check if the CRS is set, if not, set it manually
    if gdf.crs is None:
        gdf = gdf.set_crs("EPSG:4326")  # Assuming WGS84
    
    # Generate list of year-month combinations
    periods = []
    current = start
    while current <= end:
        periods.append((current.year, current.month))
        # Move to next month
        if current.month == 12:
            current = current.replace(year=current.year+1, month=1)
        else:
            current = current.replace(month=current.month+1)
    
    # Create a list to store results for each period
    all_results = []
    period_gdfs = []  # Store period GDFs for subplot generation
    
    print(f"Processing rainfall data from {start_date} to {end_date}...")
    
    # Process each time period
    for year, month in periods:
        try:
            # Define the link for CHIRPS data
            link = f"https://data.chc.ucsb.edu/products/CHIRPS-2.0/africa_monthly/tifs/chirps-v2.0.{year}.{month:02d}.tif.gz"
            
            # Download and process CHIRPS data
            with tempfile.TemporaryDirectory() as tmpdir:
                # Download the .tif.gz file
                print(f"Downloading data for {year}-{month:02d}...")
                response = requests.get(link)
                if response.status_code != 200:
                    print(f"Warning: Could not download data for {year}-{month:02d}")
                    continue
                
                zipped_file_path = os.path.join(tmpdir, "chirps.tif.gz")
                unzipped_file_path = os.path.join(tmpdir, "chirps.tif")
                
                with open(zipped_file_path, "wb") as f:
                    f.write(response.content)
                
                # Unzip the file
                with gzip.open(zipped_file_path, "rb") as f_in:
                    with open(unzipped_file_path, "wb") as f_out:
                        shutil.copyfileobj(f_in, f_out)
                
                # Open the unzipped .tif file with Rasterio
                with rasterio.open(unzipped_file_path) as src:
                    # Create a copy of the GeoDataFrame for this period
                    period_gdf = gdf.copy()
                    
                    # Reproject shapefile to match CHIRPS data CRS
                    period_gdf = period_gdf.to_crs(src.crs)
                    
                    # Calculate mean rainfall for each geometry
                    mean_rains = []
                    for geom in period_gdf.geometry:
                        try:
                            masked_data, _ = rasterio.mask.mask(src, [geom], crop=True)
                            masked_data = masked_data.flatten()
                            masked_data = masked_data[masked_data != src.nodata]  # Exclude nodata values
                            if len(masked_data) > 0:
                                mean_rains.append(masked_data.mean())
                            else:
                                mean_rains.append(np.nan)
                        except Exception as e:
                            print(f"Error processing geometry: {e}")
                            mean_rains.append(np.nan)
                    
                    # Add mean rainfall to the GeoDataFrame
                    period_gdf['mean_rain'] = mean_rains
                    period_gdf['year'] = year
                    period_gdf['month'] = month
                    period_gdf['date'] = f"{year}-{month:02d}"
                    
                    # Add to results
                    all_results.append(period_gdf)
                    period_gdfs.append((year, month, period_gdf.copy()))
                    
                    # Generate individual map for this period
                    print(f"Generating map for {year}-{month:02d}...")
                    fig, ax = plt.subplots(1, 1, figsize=(10, 10))
                    
                    # Plot the GeoDataFrame
                    period_gdf.plot(column='mean_rain', ax=ax, legend=True, 
                                   cmap='viridis', edgecolor="black", 
                                   legend_kwds={'shrink': 0.5})
                    
                    # Remove axis boxes
                    ax.set_axis_off()
                    
                    # Add title
                    plt.title(f"Mean Rainfall for {year}-{month:02d}", fontsize=16)
                    
                    # Save the map
                    map_filename = f"rainfall_{year}_{month:02d}.png"
                    map_path = maps_dir / map_filename
                    plt.savefig(map_path, bbox_inches='tight', dpi=300)
                    plt.close(fig)
                    
                    print(f"Processed {year}-{month:02d}")
                    
        except Exception as e:
            print(f"Error processing {year}-{month:02d}: {str(e)}")
    
    if not all_results:
        raise ValueError("No data could be processed for the specified date range")
    
    # Combine all results
    result_df = pd.concat(all_results)
    
    # Save data to Excel
    excel_path = output_dir / f"rainfall_data_{start_date}_to_{end_date}.xlsx"
    csv_path = output_dir / f"rainfall_data_{start_date}_to_{end_date}.csv"
    
    # Convert to regular DataFrame for Excel export (drop geometry column)
    export_df = pd.DataFrame(result_df.drop(columns=['geometry']))
    
    print(f"Saving data to Excel: {excel_path}")
    export_df.to_excel(excel_path, index=False)
    
    print(f"Saving data to CSV: {csv_path}")
    export_df.to_csv(csv_path, index=False)
    
    # Generate composite subplots figure
    print("Generating composite map with subplots (each with individual legend)...")
    create_subplots_figure(period_gdfs, maps_dir, f"rainfall_composite_{start_date}_to_{end_date}.png")
    
    print(f"Processing complete. Results saved to {output_dir}")
    print(f"- Excel data: {excel_path}")
    print(f"- CSV data: {csv_path}")
    print(f"- Maps: {maps_dir}")
    print(f"- Composite map: {maps_dir}/rainfall_composite_{start_date}_to_{end_date}.png")
    
    return result_df


def create_subplots_figure(period_gdfs, output_dir, filename):
    """
    Create a composite figure with 4 columns and automatically calculated rows
    for all time periods, with individual legends for each subplot.
    
    Parameters:
    -----------
    period_gdfs : list
        List of tuples (year, month, gdf) for each time period
    output_dir : Path
        Directory to save the figure
    filename : str
        Filename for the composite figure
    """
    n_periods = len(period_gdfs)
    if n_periods == 0:
        return
    
    # Calculate number of rows needed (4 columns)
    n_cols = 4
    n_rows = math.ceil(n_periods / n_cols)
    
    # Create figure with more space between subplots for legends
    fig_width = 24  # inches, increased for legends
    fig_height = 6 * n_rows  # 6 inches per row, increased for legends
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(fig_width, fig_height))
    
    # Flatten axes for easy indexing
    if n_rows == 1 and n_cols == 1:
        axes = np.array([axes])
    elif n_rows == 1:
        axes = np.array([axes])
    axes_flat = axes.flatten()
    
    # Plot each period
    for i, (year, month, gdf) in enumerate(period_gdfs):
        if i < len(axes_flat):
            ax = axes_flat[i]
            
            # Plot the GeoDataFrame with individual legend
            gdf.plot(column='mean_rain', ax=ax, legend=True,
                    cmap='viridis', edgecolor="black",
                    legend_kwds={'shrink': 0.5, 'aspect': 10})
            
            # Remove axis boxes
            ax.set_axis_off()
            
            # Add title
            ax.set_title(f"{year}-{month:02d}", fontsize=14)
    
    # Hide any unused subplots
    for j in range(i + 1, len(axes_flat)):
        axes_flat[j].set_visible(False)
    
    # Add a main title for the entire figure
    fig.suptitle('Rainfall by Month', fontsize=20, y=0.98)
    
    # Adjust layout
    plt.tight_layout(rect=[0, 0, 1, 0.95])  # [left, bottom, right, top]
    
    # Save the figure
    output_path = output_dir / filename
    plt.savefig(output_path, bbox_inches='tight', dpi=300)
    plt.close(fig)
    
    print(f"Composite figure saved to: {output_path}")

